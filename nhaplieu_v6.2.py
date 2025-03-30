import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog, Toplevel, Text
import pandas as pd
from datetime import datetime
import json
import os
import sys
import openpyxl
from docxtpl import DocxTemplate
from docx2pdf import convert
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import unicodedata
from PIL import Image, ImageTk
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from docxcompose.composer import Composer
from tkinterdnd2 import TkinterDnD, DND_FILES
import logging
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, Inches

from utils import create_centered_popup, add_section_break, number_to_words, normalize_vietnamese

# Thiết lập thư mục AppData và logging
appdata_dir = "AppData"
if not os.path.exists(appdata_dir):
    os.makedirs(appdata_dir)
logging.basicConfig(filename=os.path.join(appdata_dir, "app.log"), level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logging.getLogger("docxtpl").setLevel(logging.ERROR)

class ToolTip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tooltip = None
        self.widget.bind("<Enter>", self.show_tooltip)
        self.widget.bind("<Leave>", self.hide_tooltip)

    def show_tooltip(self, event):
        if self.tooltip:
            self.tooltip.destroy()
        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 5
        screen_width = self.widget.winfo_screenwidth()
        screen_height = self.widget.winfo_screenheight()
        if x + 200 > screen_width:
            x = self.widget.winfo_rootx() - 200
        if y + 30 > screen_height:
            y = self.widget.winfo_rooty() - 30
        self.tooltip = tk.Toplevel(self.widget)
        self.tooltip.wm_overrideredirect(True)
        self.tooltip.wm_geometry(f"+{x}+{y}")
        label = ttk.Label(self.tooltip, text=self.text, background="#ffffe0", padding=2, relief="solid", wraplength=200)
        label.pack()
        self.tooltip.transient(self.widget)
        self.tooltip.lift()
        self.tooltip.attributes('-topmost', True)

    def hide_tooltip(self, event):
        if self.tooltip:
            self.tooltip.destroy()
            self.tooltip = None

                
class DataEntryApp:
    def __init__(self, root):
        """Khởi tạo ứng dụng nhập liệu hồ sơ kinh doanh."""
        self.root = root
        self.root.title("Ứng dụng nhập liệu hồ sơ Kinh Doanh v6.0")
        window_width = 1600
        window_height = 800
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")
        #self.root.state('zoomed')  # Tùy chọn: Phóng to tối đa khi khởi động       
        self.root.configure(bg="#f8f9fa")
        logging.info("Ứng dụng khởi động")

        # Cấu hình file trong thư mục AppData
        self.appdata_dir = appdata_dir
        self.configs_file = os.path.join(self.appdata_dir, "configs.json")
        self.templates_dir = "templates"
        self.backup_dir = "backup"

        if not os.path.exists(self.templates_dir):
            os.makedirs(self.templates_dir)
        if not os.path.exists(self.backup_dir):
            os.makedirs(self.backup_dir)

        self.default_fields = [
            "ten_doanh_nghiep", "ten_doanh_nghiep_bang_tieng_anh", "ten_viet_tat",
            "ma_so_doanh_nghiep", "ngay_cap_mst", "so_nha_ten_duong", "xa_phuong", "quan_huyen",
            "tinh_thanh_pho", "so_dien_thoai", "von_đieu_le", "ho_ten", "gioi_tinh", "sinh_ngay",
            "dan_toc", "quoc_tich", "so_cccd", "ngay_cap", "noi_cap", "ngay_het_han", "dia_chi_thuong_tru", "dia_chi_lien_lac",
            "ho_ten_uq", "gioi_tinh_uq", "sinh_ngay_uq", "so_cccd_uq", "ngay_cap_uq", "noi_cap_uq", "dia_chi_lien_lac_uq"
        ]
        # Định nghĩa các cột cho bảng thông tin thành viên
        self.member_columns = [
            "ho_ten", "gioi_tinh", "ngay_sinh", "dan_toc", "quoc_tich", "loai_giay_to",
            "so_cccd", "ngay_cap", "noi_cap",
            "ngay_het_han", "dia_chi_thuong_tru", "dia_chi_lien_lac", "von_gop", "ty_le_gop", "ngay_gop_von"
        ]

        self.configs = {}
        self.current_config_name = None
        self.field_groups = {}
        self.fields = []
        self.saved_entries = []
        self.drag_item = None
        self.entries = {}
        self.labels = {}
        self.current_tab_index = 0
#================================ KHU VỰC QUẢN LÝ GIAO DIỆN ========================================
        # Khởi tạo top_frame trước
        self.top_frame = ttk.Frame(self.root)
        self.top_frame.pack(side="top", fill="x", padx=5, pady=5)
        
        self.main_frame = ttk.Frame(root, padding=10)
        self.main_frame.pack(fill="both", expand=True)

        # Config Frame
        self.config_frame = ttk.LabelFrame(self.main_frame, text="Quản lý cấu hình", padding=10)
        self.config_frame.pack(fill="x", pady=(0, 10))

        self.config_var = tk.StringVar()
        self.config_dropdown = ttk.Combobox(self.config_frame, textvariable=self.config_var, 
                                           values=list(self.configs.keys()), state="readonly", width=15)
        self.config_dropdown.pack(side="left", padx=5)
        self.config_dropdown.bind("<<ComboboxSelected>>", self.load_selected_config)

        def resource_path(relative_path): # Hàm hỗ trợ lấy đường dẫn tài nguyên
            """Lấy đường dẫn tuyệt đối đến tài nguyên, hoạt động cho cả mã nguồn và file thực thi PyInstaller."""
            if hasattr(sys, '_MEIPASS'):
                return os.path.join(sys._MEIPASS, relative_path)
            return os.path.join(os.path.abspath("."), relative_path)

        # Load icons
        try:
            add_icon = Image.open(resource_path("icon/add_icon.png")).resize((20, 20), Image.Resampling.LANCZOS)
            delete_icon = Image.open(resource_path("icon/delete_icon.png")).resize((20, 20), Image.Resampling.LANCZOS)
            save_icon = Image.open(resource_path("icon/save_icon.png")).resize((20, 20), Image.Resampling.LANCZOS)
            clear_icon = Image.open(resource_path("icon/clear_icon.png")).resize((20, 20), Image.Resampling.LANCZOS)
            import_icon = Image.open(resource_path("icon/imxls_icon.png")).resize((20, 20), Image.Resampling.LANCZOS)
            xls_icon = Image.open(resource_path("icon/xls_icon.png")).resize((20, 20), Image.Resampling.LANCZOS)
            preview_icon = Image.open(resource_path("icon/preview_icon.png")).resize((40, 40), Image.Resampling.LANCZOS)
            export_icon = Image.open(resource_path("icon/export_icon.png")).resize((40, 40), Image.Resampling.LANCZOS)
            remove_icon = Image.open(resource_path("icon/remove_icon.png")).resize((20, 20), Image.Resampling.LANCZOS)
            edit_icon = Image.open(resource_path("icon/edit_icon.png")).resize((20, 20), Image.Resampling.LANCZOS)
            search_icon = Image.open(resource_path("icon/search_icon.png")).resize((20, 20), Image.Resampling.LANCZOS)
            restorebackup_icon = Image.open(resource_path("icon/restorebackup_icon.png")).resize((20, 20), Image.Resampling.LANCZOS)
            self.add_icon_img = ImageTk.PhotoImage(add_icon)
            self.delete_icon_img = ImageTk.PhotoImage(delete_icon)
            self.save_icon_img = ImageTk.PhotoImage(save_icon)
            self.clear_icon_img = ImageTk.PhotoImage(clear_icon)
            self.import_icon_img = ImageTk.PhotoImage(import_icon)
            self.xls_icon_img = ImageTk.PhotoImage(xls_icon)
            self.preview_icon_img = ImageTk.PhotoImage(preview_icon)
            self.export_icon_img = ImageTk.PhotoImage(export_icon)
            self.remove_icon_img = ImageTk.PhotoImage(remove_icon)
            self.edit_icon_img = ImageTk.PhotoImage(edit_icon)
            self.search_icon_img = ImageTk.PhotoImage(search_icon)
            self.restorebackup_icon_img = ImageTk.PhotoImage(restorebackup_icon)
        except Exception as e:
            logging.error(f"Lỗi khi tải biểu tượng: {str(e)}")
            messagebox.showerror("Lỗi", "Không thể tải biểu tượng, kiểm tra thư mục icon!")

        #Quản lý cấu hình
        self.add_config_button = ttk.Button(self.config_frame, image=self.add_icon_img, command=self.add_new_config, style="info.TButton")
        ToolTip(self.add_config_button, "Thêm cấu hình")
        self.add_config_button.pack(side="left", padx=5)
        self.delete_config_button = ttk.Button(self.config_frame, image=self.delete_icon_img, command=self.delete_config, style="danger.TButton")
        ToolTip(self.delete_config_button, "Xóa cấu hình")
        self.delete_config_button.pack(side="left", padx=5)
        self.rename_config_button = ttk.Button(self.config_frame, image=self.edit_icon_img, command=self.rename_config, style="warning.TButton")
        ToolTip(self.rename_config_button, "Sửa tên cấu hình")
        self.rename_config_button.pack(side="left", padx=5)

        
        #Quản lý tab
        ttk.Separator(self.config_frame, orient="vertical").pack(side="left", padx=10, fill="y")
        ttk.Label(self.config_frame, text="Quản lý tab:").pack(side="left", padx=5)
        self.tab_var = tk.StringVar()
        self.tab_dropdown = ttk.Combobox(self.config_frame, textvariable=self.tab_var, state="readonly", width=20)
        self.tab_dropdown.pack(side="left", padx=5)
        self.add_tab_button = ttk.Button(self.config_frame, image=self.add_icon_img, command=self.add_tab, style="info.TButton")
        ToolTip(self.add_tab_button, "Thêm tab")
        self.add_tab_button.pack(side="left", padx=5)
        self.delete_tab_button = ttk.Button(self.config_frame, image=self.delete_icon_img, command=self.delete_tab, style="danger.TButton")
        ToolTip(self.delete_tab_button, "Xóa tab")
        self.delete_tab_button.pack(side="left", padx=5)
        self.rename_tab_button = ttk.Button(self.config_frame, image=self.edit_icon_img, command=self.rename_tab, style="warning.TButton")
        ToolTip(self.rename_tab_button, "Sửa tên tab")
        self.rename_tab_button.pack(side="left", padx=5)

        #Quản lý trường
        ttk.Separator(self.config_frame, orient="vertical").pack(side="left", padx=10, fill="y")
        ttk.Label(self.config_frame, text="Quản lý trường:").pack(side="left", padx=5)
        self.field_var = tk.StringVar()
        self.field_dropdown = ttk.Combobox(self.config_frame, textvariable=self.field_var, state="readonly", width=20)
        self.field_dropdown.pack(side="left", padx=5)
        # Thêm nút "Thêm trường"
        self.add_field_button = ttk.Button(self.config_frame, image=self.add_icon_img, command=self.add_field, style="info.TButton")
        ToolTip(self.add_field_button, "Thêm trường")
        self.add_field_button.pack(side="left", padx=5)
        self.delete_field_button = ttk.Button(self.config_frame, image=self.delete_icon_img, command=self.delete_selected_field, style="danger.TButton")
        ToolTip(self.delete_field_button, "Xóa trường")
        self.delete_field_button.pack(side="left", padx=5)
        self.rename_field_button = ttk.Button(self.config_frame, image=self.edit_icon_img, command=self.rename_selected_field, style="warning.TButton")
        ToolTip(self.rename_field_button, "Sửa tên trường")
        self.rename_field_button.pack(side="left", padx=5)

        #Quản lý template"
        ttk.Separator(self.config_frame, orient="vertical").pack(side="left", padx=10, fill="y")
        ttk.Label(self.config_frame, text="Quản lý mẫu:").pack(side="left", padx=5) 
        self.template_frame = ttk.Frame(self.config_frame)
        self.template_frame.pack(side="left", fill="x", padx=5)
        self.template_tree = ttk.Treeview(self.template_frame, columns=(), show="tree", height=5, selectmode="extended")
        self.template_tree.column("#0", width=200)
        self.template_tree.pack(fill="both", expand=True)
        self.update_template_tree()
        self.template_frame.drop_target_register(DND_FILES)
        self.template_frame.dnd_bind('<<Drop>>', self.drop_template_files)
        ToolTip(self.template_frame, "Kéo thả để thêm/sắp xếp")
        self.template_tree.bind("<Button-1>", self.start_drag)
        self.template_tree.bind("<B1-Motion>", self.drag_template)
        self.template_tree.bind("<ButtonRelease-1>", self.drop_template)
        self.template_tree.bind("<Button-3>", self.show_template_context_menu)


        # Control Frame
        self.control_frame = ttk.LabelFrame(self.main_frame, text="Quản lý dữ liệu", padding=10)
        self.control_frame.pack(fill="x", pady=10)

        ttk.Label(self.control_frame, text="Tên dữ liệu:").grid(row=0, column=0, padx=5, pady=5)
        self.load_data_var = tk.StringVar()
        self.load_data_dropdown = ttk.Combobox(self.control_frame, textvariable=self.load_data_var, state="readonly", width=40)
        self.load_data_dropdown.grid(row=0, column=1, padx=5, pady=5)
        self.load_data_dropdown.bind("<<ComboboxSelected>>", self.load_selected_entry)
        self.add_data_button = ttk.Button(self.control_frame, image=self.add_icon_img, command=self.add_entry_data, style="info.TButton")
        ToolTip(self.add_data_button, "Thêm dữ liệu")
        self.add_data_button.grid(row=0, column=2, padx=5, pady=5)
        self.delete_data_button = ttk.Button(self.control_frame, image=self.delete_icon_img, command=self.delete_entry_data, style="danger.TButton")
        ToolTip(self.delete_data_button, "Xóa dữ liệu")
        self.delete_data_button.grid(row=0, column=3, padx=5, pady=5)
        self.rename_data_button = ttk.Button(self.control_frame, image=self.edit_icon_img, command=self.rename_entry_data, style="warning.TButton")
        ToolTip(self.rename_data_button, "Sửa tên dữ liệu")
        self.rename_data_button.grid(row=0, column=4, padx=5, pady=5)

        ttk.Separator(self.control_frame, orient="vertical").grid(row=0, column=5, padx=10, pady=5, sticky="ns")
        self.edit_data_button = ttk.Button(self.control_frame, image=self.save_icon_img, command=self.save_entry_data, style="success.TButton")
        ToolTip(self.edit_data_button, "Lưu thông tin")
        self.edit_data_button.grid(row=0, column=6, padx=5, pady=5)
        self.clear_data_button = ttk.Button(self.control_frame, image=self.clear_icon_img, command=self.clear_entries, style="danger.TButton")
        ToolTip(self.clear_data_button, "Xóa thông tin")
        self.clear_data_button.grid(row=0, column=7, padx=5, pady=5)

        ttk.Separator(self.control_frame, orient="vertical").grid(row=0, column=8, padx=10, pady=5, sticky="ns")
        self.export_excel_button = ttk.Button(self.control_frame, image=self.xls_icon_img, command=self.export_data, style="primary.TButton")
        ToolTip(self.export_excel_button, "Xuất dữ liệu")
        self.export_excel_button.grid(row=0, column=9, padx=5, pady=5)
        self.import_data_button = ttk.Button(self.control_frame, image=self.import_icon_img, command=self.import_from_file, style="primary.TButton")
        ToolTip(self.import_data_button, "Nhập dữ liệu")
        self.import_data_button.grid(row=0, column=10, padx=5, pady=5)
        
        self.restore_data_button = ttk.Button(self.control_frame, image=self.restorebackup_icon_img, command=self.restore_from_backup, style="primary.TButton")
        ToolTip(self.restore_data_button, "Khôi phục từ sao lưu")
        self.restore_data_button.grid(row=0, column=11, padx=5, pady=5)

        ttk.Separator(self.control_frame, orient="vertical").grid(row=0, column=12, padx=10, pady=5, sticky="ns")
        self.show_placeholder_button = ttk.Button(self.control_frame, image=self.search_icon_img, command=self.show_placeholder_popup, style="danger.TButton")
        ToolTip(self.show_placeholder_button, "Hiển thị danh sách placeholder")
        self.show_placeholder_button.grid(row=0, column=13, padx=5, pady=5)

        # Notebook với tiêu đề "Quản lý nhập liệu"
        self.notebook_frame = ttk.LabelFrame(self.main_frame, text="Quản lý nhập liệu", padding=5)
        self.notebook_frame.pack(fill="both", expand=True, pady=10)
        self.notebook = ttk.Notebook(self.notebook_frame)
        self.notebook.pack(fill="both", expand=True)
        # Gắn sự kiện khi tab thay đổi
        self.notebook.bind("<<NotebookTabChanged>>", lambda event: self.update_field_dropdown())

        
        # Export Frame
        self.export_frame = ttk.Frame(self.main_frame)
        self.export_frame.pack(fill="x", pady=10)
        self.export_frame.grid_columnconfigure(0, weight=1)
        self.export_frame.grid_columnconfigure(3, weight=1)
        self.preview_word_button = ttk.Button(self.export_frame, image=self.preview_icon_img, command=self.preview_word, style="secondary.TButton")
        ToolTip(self.preview_word_button, "Xem trước")
        self.preview_word_button.grid(row=0, column=1, padx=5)
        self.export_file_button = ttk.Button(self.export_frame, image=self.export_icon_img, command=self.export_file, style="danger.TButton")
        ToolTip(self.export_file_button, "Xuất file")
        self.export_file_button.grid(row=0, column=2, padx=5)

        # Styles
        style = ttk.Style()
        style.theme_use("flatly")
        style.configure("TButton", font=("Segoe UI", 10), borderwidth=1, relief="flat", padding=5)
        style.configure("info.TButton", background="#4582ec", foreground="white", bordercolor="#4582ec", borderradius=90)
        style.map("info.TButton", background=[("active", "#3572d8")])
        style.configure("danger.TButton", background="#adb5bd", foreground="black", bordercolor="#adb5bd", borderradius=90, focusthickness=0)
        style.map("danger.TButton", background=[("active", "#9ba3ab")], bordercolor=[("active", "#9ba3ab"), ("focus", "#adb5bd")], foreground=[("active", "black")])
        style.configure("warning.TButton", background="#f0ad4e", foreground="black", bordercolor="#f0ad4e", borderradius=90)
        style.map("warning.TButton", background=[("active", "#e89b3c")])
        style.configure("success.TButton", background="#02b875", foreground="white", bordercolor="#02b875", borderradius=90)
        style.map("success.TButton", background=[("active", "#019c62")])
        style.configure("primary.TButton", background="#4582ec", foreground="white", bordercolor="#4582ec", borderradius=90)
        style.map("primary.TButton", background=[("active", "#3572d8")])
        style.configure("secondary.TButton", background="#4582ec", foreground="white", bordercolor="#4582ec", borderradius=90)
        style.map("secondary.TButton", background=[("active", "#3572d8")])
        style.configure("TLabel", font=("Segoe UI", 10))
        style.configure("TEntry", font=("Segoe UI", 10))
        style.configure("TCombobox", font=("Segoe UI", 10))

        self.load_configs()
        if not self.configs:
            self.initialize_default_config()
        self.config_dropdown["values"] = list(self.configs.keys())
        if self.configs:
            self.config_dropdown.set(list(self.configs.keys())[0])
            self.load_selected_config(None)
        
        self.root.after(600000, self.auto_backup)
     
    
    
        
#================================ LOAD CẤU HÌNH ========================================

    def initialize_default_config(self):
        """Khởi tạo cấu hình mặc định nếu không có file configs."""
        self.configs = {
            "Mặc định": {
                "field_groups": {
                    "Thông tin công ty": self.default_fields[0:10],
                    "Thông tin ĐDPL": self.default_fields[10:20],
                    "Thông tin thành viên": [],  # Thêm tab cho thành viên
                    "Thông tin uỷ quyền": self.default_fields[20:],
                    "Ngành nghề kinh doanh": []
                    
                },
                "templates": {},
                "entries": []
            }
        }
        self.save_configs()
        self.current_config_name = "Mặc định"
        self.field_groups = self.configs["Mặc định"]["field_groups"]
        self.saved_entries = self.configs["Mặc định"]["entries"]
        self.fields = self.default_fields.copy()
        
    def load_configs(self):
        """Tải cấu hình từ file JSON trong AppData."""
        try:
            if os.path.exists(self.configs_file):
                with open(self.configs_file, 'r', encoding='utf-8') as f:
                    self.configs = json.load(f)
            else:
                self.initialize_default_config()
        except Exception as e:
            logging.error(f"Lỗi khi tải configs: {str(e)}")
            messagebox.showerror("Lỗi", f"Không thể tải cấu hình: {str(e)}")

    def load_selected_config(self, event):
        """Tải cấu hình được chọn từ dropdown."""
        self.current_config_name = self.config_var.get()
        self.field_groups = self.configs.get(self.current_config_name, {}).get("field_groups", {})
        self.saved_entries = self.configs.get(self.current_config_name, {}).get("entries", [])  
        '''#Đảm bảo tab "Ngành nghề kinh doanh" luôn tồn tại    
        if "Ngành nghề kinh doanh" not in self.field_groups:
            #self.field_groups["Ngành nghề kinh doanh"] = []
            #self.configs[self.current_config_name]["field_groups"] = self.field_groups
            #self.save_configs()'''
        
        self.fields = [field for fields in self.field_groups.values() for field in fields if fields]
        self.clear_tabs()
        self.create_tabs()
        self.tab_dropdown["values"] = list(self.field_groups.keys())
        self.tab_dropdown.set(list(self.field_groups.keys())[0] if self.field_groups else "")
        
        self.load_data_dropdown.set("")
        self.load_data_dropdown["values"] = [entry["name"] for entry in self.saved_entries]
        self.update_template_tree()
        
        if self.saved_entries:
            self.load_data_dropdown.set(self.saved_entries[0]["name"])
            self.load_selected_entry(None)
            
        self.update_field_dropdown()  # Cập nhật danh sách trường khi tải cấu hình

    def load_selected_entry(self, event):
        selected_name = self.load_data_var.get()
        for entry in self.saved_entries:
            if entry["name"] == selected_name:
                for field, value in entry["data"].items():
                    if field in self.entries:
                        self.entries[field].delete(0, tk.END)
                        self.entries[field].insert(0, value)
                if hasattr(self, 'industry_tree'): # Kiểm tra xem self.load_industry_data() có tồn tại không trước khi gọi load_industry_data
                    self.load_industry_data()
                if hasattr(self, 'member_tree'):
                    self.load_member_data()  # Kiểm tra xem self.member_tree có tồn tại không trước khi gọi load_member_data
                break

    def update_field_dropdown(self):
        current_tab = self.notebook.tab(self.notebook.select(), "text")
        fields = self.field_groups.get(current_tab, [])
        self.field_dropdown["values"] = fields
        self.field_var.set(fields[0] if fields else "")
        
#================================ KHU VỰC QUẢN LÝ CẤU HÌNH ========================================
    def add_new_config(self):
        config_name = simpledialog.askstring("Thêm cấu hình", "Nhập tên cấu hình mới:")
        if config_name and config_name not in self.configs:
            self.configs[config_name] = {
                "field_groups": {
                    "Thông tin công ty": [],
                    "Thông tin ĐDPL": [],
                    "Thông tin thành viên": [],
                    "Thông tin uỷ quyền": [],
                    "Ngành nghề kinh doanh": []  
                },
                "templates": {},
                "entries": []
            }
            self.save_configs()
            self.config_dropdown["values"] = list(self.configs.keys())
            self.config_dropdown.set(config_name)
            self.load_selected_config(None)
            logging.info(f"Thêm cấu hình '{config_name}'")

    def delete_config(self):
        if not self.current_config_name:
            return
        if messagebox.askyesno("Xác nhận", f"Bạn có chắc muốn xóa cấu hình '{self.current_config_name}' không?"):
            del self.configs[self.current_config_name]
            self.save_configs()
            self.config_dropdown["values"] = list(self.configs.keys())
            if self.configs:
                self.config_dropdown.set(list(self.configs.keys())[0])
                self.load_selected_config(None)
            else:
                self.initialize_default_config()
                self.config_dropdown.set("Mặc định")
                self.load_selected_config(None)
            logging.info(f"Xóa cấu hình '{self.current_config_name}'")

    def rename_config(self):
        old_name = self.current_config_name
        if not old_name:
            return
        new_name = simpledialog.askstring("Sửa tên cấu hình", "Nhập tên mới:", initialvalue=old_name)
        if new_name and new_name != old_name and new_name not in self.configs:
            self.configs[new_name] = self.configs.pop(old_name)
            self.current_config_name = new_name
            self.save_configs()
            self.config_dropdown["values"] = list(self.configs.keys())
            self.config_dropdown.set(new_name)
            messagebox.showinfo("Thành công", f"Đã đổi tên thành '{new_name}'!")
            logging.info(f"Đổi tên cấu hình từ '{old_name}' thành '{new_name}'")

    def save_configs(self):
        """Lưu cấu hình (bao gồm field_groups, templates, và entries) vào file JSON trong AppData."""
        try:
            with open(self.configs_file, 'w', encoding='utf-8') as f:
                json.dump(self.configs, f, ensure_ascii=False, indent=4)
            logging.info("Đã lưu cấu hình")
        except Exception as e:
            logging.error(f"Lỗi khi lưu configs: {str(e)}")
            messagebox.showerror("Lỗi", f"Không thể lưu cấu hình: {str(e)}")

#================================ KHU VỰC QUẢN LÝ TRƯỜNG ========================================
    def add_field(self):
        if not self.current_config_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn cấu hình trước!")
            return
        new_field = simpledialog.askstring("Thêm trường", "Nhập tên trường mới:")
        if new_field:
            new_field = new_field.strip()
            if not new_field:
                messagebox.showwarning("Cảnh báo", "Tên trường không được để trống!")
                return
            current_tab = self.notebook.tab(self.notebook.select(), "text")  # Lấy tab hiện tại từ notebook
            if not current_tab:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn tab trước khi thêm trường!")
                return
            if new_field in self.field_groups[current_tab]:
                messagebox.showwarning("Cảnh báo", "Trường này đã tồn tại trong tab hiện tại!")
                return
            self.field_groups[current_tab].append(new_field)
            self.configs[self.current_config_name]["field_groups"][current_tab] = self.field_groups[current_tab]
            self.save_configs()
            self.load_selected_config(None)
            logging.info(f"Thêm trường '{new_field}' vào tab '{current_tab}'")
            
    def delete_selected_field(self):
        if not self.current_config_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn cấu hình trước!")
            return
        field = self.field_var.get()
        if not field:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn trường để xóa!")
            return
        current_tab = self.notebook.tab(self.notebook.select(), "text")
        if field in self.field_groups[current_tab]:
            self.field_groups[current_tab].remove(field)
            self.configs[self.current_config_name]["field_groups"][current_tab] = self.field_groups[current_tab]
            self.save_configs()
            self.load_selected_config(None)
            logging.info(f"Xóa trường '{field}' khỏi tab '{current_tab}'")

    def rename_selected_field(self):
        if not self.current_config_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn cấu hình trước!")
            return
        old_field = self.field_var.get()
        if not old_field:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn trường để sửa tên!")
            return
        new_field = simpledialog.askstring("Sửa tên trường", "Nhập tên mới cho trường:", initialvalue=old_field)
        if new_field and new_field != old_field:
            current_tab = self.notebook.tab(self.notebook.select(), "text")
            if new_field in self.field_groups[current_tab]:
                messagebox.showwarning("Cảnh báo", "Tên trường đã tồn tại trong tab này!")
                return
            idx = self.field_groups[current_tab].index(old_field)
            self.field_groups[current_tab][idx] = new_field
            self.configs[self.current_config_name]["field_groups"][current_tab] = self.field_groups[current_tab]
            self.save_configs()
            self.load_selected_config(None)
            logging.info(f"Sửa tên trường từ '{old_field}' thành '{new_field}' trong tab '{current_tab}'")

    def rename_field(self, field):
        if not self.current_config_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn cấu hình trước!")
            return
        old_field = field
        new_field = simpledialog.askstring("Sửa tên trường", "Nhập tên mới cho trường:", initialvalue=old_field)
        if new_field and new_field != old_field:
            current_tab = self.notebook.tab(self.notebook.select(), "text")
            if new_field in self.field_groups[current_tab]:
                messagebox.showwarning("Cảnh báo", "Tên trường đã tồn tại trong tab này!")
                return
            idx = self.field_groups[current_tab].index(old_field)
            self.field_groups[current_tab][idx] = new_field
            self.configs[self.current_config_name]["field_groups"][current_tab] = self.field_groups[current_tab]
            self.save_configs()
            self.load_selected_config(None)
            logging.info(f"Sửa tên trường từ '{old_field}' thành '{new_field}' trong tab '{current_tab}'")

    def delete_field(self, field):
        if not self.current_config_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn cấu hình trước!")
            return
        current_tab = self.notebook.tab(self.notebook.select(), "text")
        if field in self.field_groups[current_tab]:
            self.field_groups[current_tab].remove(field)
            self.configs[self.current_config_name]["field_groups"][current_tab] = self.field_groups[current_tab]
            self.save_configs()
            self.load_selected_config(None)
            logging.info(f"Xóa trường '{field}' khỏi tab '{current_tab}'")
            
#================================ KHU VỰC QUẢN LÝ DỮ LIỆU ========================================

    def add_entry_data(self):
        if not self.current_config_name:
            return
        entry_name = simpledialog.askstring("Thêm dữ liệu", "Nhập tên cho bộ dữ liệu mới:")
        if entry_name:
            data = {field: "" for field in self.fields}
            data["nganh_nghe"] = [] 
            self.saved_entries.append({"name": entry_name, "data": data})
            self.configs[self.current_config_name]["entries"] = self.saved_entries
            self.load_data_dropdown["values"] = [entry["name"] for entry in self.saved_entries]
            self.save_configs()
            self.load_data_dropdown.set(entry_name)
            self.load_selected_entry(None)
            messagebox.showinfo("Thành công", f"Đã thêm dữ liệu '{entry_name}'!")
            logging.info(f"Thêm dữ liệu '{entry_name}'")

    def save_entry_data(self):
        if not self.current_config_name:
            return
        data = {field: self.entries[field].get() for field in self.entries}
        selected_name = self.load_data_var.get()
        if not selected_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn dữ liệu để lưu!")
            return
        
        # Tự động thêm von_đieu_le_bang_chu vào dữ liệu
        if "von_đieu_le" in data:
            data["von_đieu_le_bang_chu"] = number_to_words(data["von_đieu_le"])
        
        # Lấy danh sách ngành nghề từ industry_tree
        industries = []
        if hasattr(self, 'industry_tree'):
            for item in self.industry_tree.get_children():
                values = self.industry_tree.item(item)['values']
                industry = {
                    "ten_nganh": values[0],
                    "ma_nganh": values[1],
                    "la_nganh_chinh": values[2] == "X"
                }
                industries.append(industry)
        
        for i, entry in enumerate(self.saved_entries):
            if entry["name"] == selected_name:
                entry["data"].update(data)
                entry["data"]["nganh_nghe"] = industries  # Cập nhật danh sách ngành nghề
                self.configs[self.current_config_name]["entries"] = self.saved_entries
                self.save_configs()
                messagebox.showinfo("Thành công", f"Đã lưu dữ liệu '{selected_name}'!")
                logging.info(f"Lưu dữ liệu '{selected_name}'")
                break

    def delete_entry_data(self):
        if not self.current_config_name:
            return
        selected_name = self.load_data_var.get()
        if not selected_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn dữ liệu để xóa!")
            return
        if messagebox.askyesno("Xác nhận", f"Bạn có chắc muốn xóa dữ liệu '{selected_name}' không?"):
            self.saved_entries = [entry for entry in self.saved_entries if entry["name"] != selected_name]
            self.configs[self.current_config_name]["entries"] = self.saved_entries
            self.load_data_dropdown["values"] = [entry["name"] for entry in self.saved_entries]
            self.save_configs()
            messagebox.showinfo("Thành công", f"Dữ liệu '{selected_name}' đã được xóa!")
            logging.info(f"Xóa dữ liệu '{selected_name}'")
            for entry in self.entries.values():
                entry.delete(0, tk.END)

    def rename_entry_data(self):
        if not self.current_config_name:
            return
        selected_name = self.load_data_var.get()
        if not selected_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn dữ liệu để sửa tên!")
            return
        new_name = simpledialog.askstring("Sửa tên dữ liệu", "Nhập tên mới:", initialvalue=selected_name)
        if new_name and new_name != selected_name:
            if new_name in [entry["name"] for entry in self.saved_entries]:
                messagebox.showwarning("Cảnh báo", "Tên này đã tồn tại, vui lòng chọn tên khác!")
                return
            for entry in self.saved_entries:
                if entry["name"] == selected_name:
                    entry["name"] = new_name
                    break
            self.configs[self.current_config_name]["entries"] = self.saved_entries
            self.load_data_dropdown["values"] = [entry["name"] for entry in self.saved_entries]
            self.load_data_dropdown.set(new_name)
            self.save_configs()
            messagebox.showinfo("Thành công", f"Đã đổi tên dữ liệu thành '{new_name}'!")
            logging.info(f"Đổi tên dữ liệu từ '{selected_name}' thành '{new_name}'")

#================================ KHU VỰC QUẢN LÝ NHẬP DỮ LIỆU, XUẤT DỮ LIỆU  ========================================
    def export_data(self):
        """Xuất dữ liệu ra file Excel hoặc JSON."""
        # Tạo popup
        popup = create_centered_popup (self.root, "Chọn định dạng xuất", 300, 150)
        popup.bind('<Escape>', lambda e: popup.destroy())
        ttk.Label(popup, text="Chọn định dạng:").pack(pady=10)
        file_format = tk.StringVar(value="Excel")
        ttk.Radiobutton(popup, text="Excel (.xlsx)", variable=file_format, value="Excel").pack(anchor="w", padx=20)
        ttk.Radiobutton(popup, text="JSON (.json)", variable=file_format, value="JSON").pack(anchor="w", padx=20)

        def confirm_export():
            popup.destroy()
            if file_format.get() == "Excel":
                output_path = filedialog.asksaveasfilename(
                    defaultextension=".xlsx",
                    filetypes=[("Excel files", "*.xlsx")],
                    initialfile=f"data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                )
                if not output_path:
                    return
                data = {field: self.entries[field].get() for field in self.entries}
                selected_name = self.load_data_var.get()
                for entry in self.saved_entries:
                    if entry["name"] == selected_name:
                        industries = entry["data"].get("nganh_nghe", [])
                        members = entry["data"].get("thanh_vien", [])
                        df_main = pd.DataFrame([data])
                        df_industries = pd.DataFrame(industries)
                        df_members = pd.DataFrame(members)
                        break
                else:
                    df_main = pd.DataFrame([data])
                    df_industries = pd.DataFrame()
                    df_members = pd.DataFrame()

                try:
                    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
                        df_main.to_excel(writer, sheet_name="Main", index=False)
                        df_industries.to_excel(writer, sheet_name="Industries", index=False)
                        df_members.to_excel(writer, sheet_name="Members", index=False)
                    os.startfile(output_path)
                    messagebox.showinfo("Thành công", f"Đã xuất file Excel: {output_path}")
                    logging.info(f"Xuất Excel: {output_path}")
                except Exception as e:
                    logging.error(f"Lỗi khi xuất Excel: {str(e)}")
                    messagebox.showerror("Lỗi", f"Không thể xuất file Excel: {str(e)}")

            elif file_format.get() == "JSON":
                output_path = filedialog.asksaveasfilename(
                    defaultextension=".json",
                    filetypes=[("JSON files", "*.json")],
                    initialfile=f"data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                )
                if not output_path:
                    return
                data = {field: self.entries[field].get() for field in self.entries}
                selected_name = self.load_data_var.get()
                export_data = None
                for entry in self.saved_entries:
                    if entry["name"] == selected_name:
                        export_data = {
                            "name": selected_name,
                            "data": {
                                **data,
                                "nganh_nghe": entry["data"].get("nganh_nghe", []),
                                "thanh_vien": entry["data"].get("thanh_vien", [])
                            }
                        }
                        break
                if export_data is None:
                    export_data = {
                        "name": selected_name,
                        "data": {
                            **data,
                            "nganh_nghe": [],
                            "thanh_vien": []
                        }
                    }

                try:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        json.dump(export_data, f, ensure_ascii=False, indent=4)
                    os.startfile(output_path)
                    messagebox.showinfo("Thành công", f"Đã xuất file JSON: {output_path}")
                    logging.info(f"Xuất JSON: {output_path}")
                except Exception as e:
                    logging.error(f"Lỗi khi xuất JSON: {str(e)}")
                    messagebox.showerror("Lỗi", f"Không thể xuất file JSON: {str(e)}")

        ttk.Button(popup, text="Xuất", command=confirm_export, style="primary.TButton").pack(pady=10)

    def import_from_file(self):
        """Nhập dữ liệu từ file Excel hoặc JSON."""
        file_path = filedialog.askopenfilename(filetypes=[("JSON files", "*.json"), ("Excel files", "*.xlsx")])
        if not file_path:
            return
        try:
            if file_path.endswith(".xlsx"):
                # Đọc sheet "Main" cho dữ liệu chính
                df_main = pd.read_excel(file_path, sheet_name="Main")
                # Thay thế NaN bằng chuỗi rỗng trong dữ liệu chính
                data = df_main.iloc[0].fillna("").to_dict()
                for field in self.entries:
                    if field in data:
                        self.entries[field].delete(0, tk.END)
                        self.entries[field].insert(0, str(data[field]))

                # Đọc sheet "Industries" cho dữ liệu ngành nghề
                try:
                    df_industries = pd.read_excel(file_path, sheet_name="Industries")
                    # Thay thế NaN bằng chuỗi rỗng trong dữ liệu ngành nghề
                    industries = df_industries.fillna("").to_dict(orient="records")
                except ValueError:
                    industries = []

                # Đọc sheet "Members" cho dữ liệu thành viên
                try:
                    df_members = pd.read_excel(file_path, sheet_name="Members")
                    # Thay thế NaN bằng chuỗi rỗng trong dữ liệu thành viên
                    members = df_members.fillna("").to_dict(orient="records")
                except ValueError:
                    members = []

                selected_name = self.load_data_var.get()
                for entry in self.saved_entries:
                    if entry["name"] == selected_name:
                        entry["data"] = {field: self.entries[field].get() for field in self.entries}
                        entry["data"]["nganh_nghe"] = industries
                        entry["data"]["thanh_vien"] = members
                        break
                else:
                    new_entry = {
                        "name": selected_name,
                        "data": {field: self.entries[field].get() for field in self.entries}
                    }
                    new_entry["data"]["nganh_nghe"] = industries
                    new_entry["data"]["thanh_vien"] = members
                    self.saved_entries.append(new_entry)
                    self.load_data_dropdown["values"] = [entry["name"] for entry in self.saved_entries]

                self.save_configs()
                self.load_member_data()
                self.load_industry_data()
                messagebox.showinfo("Thành công", "Đã nhập dữ liệu từ file Excel!")
                logging.info(f"Nhập dữ liệu từ Excel: {file_path}")
            elif file_path.endswith(".json"):
                with open(file_path, 'r', encoding='utf-8') as f:
                    imported_data = json.load(f)
                if isinstance(imported_data, dict) and "name" in imported_data and "data" in imported_data:
                    self.saved_entries.append(imported_data)
                    self.load_data_dropdown["values"] = [entry["name"] for entry in self.saved_entries]
                    self.load_data_var.set(imported_data["name"])
                    self.load_selected_entry(None)
                    self.save_configs()
                    messagebox.showinfo("Thành công", "Đã nhập dữ liệu từ file JSON!")
                    logging.info(f"Nhập dữ liệu từ JSON: {file_path}")
        except Exception as e:
            logging.error(f"Lỗi khi nhập file: {str(e)}")
            messagebox.showerror("Lỗi", f"Không thể nhập file: {str(e)}")
            
    def show_placeholder_popup(self):
        """Hiển thị popup chứa danh sách placeholder và nút xuất."""
        if not self.current_config_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn cấu hình trước!")
            return

        popup = create_centered_popup(self.root, "Danh sách Placeholder", 600, 400)
        popup.bind('<Escape>', lambda e: popup.destroy())
        placeholder_frame = ttk.Frame(popup)
        placeholder_frame.pack(fill="both", expand=True, padx=10, pady=10)
        placeholder_list = ttk.Treeview(placeholder_frame, columns=("original"), show="tree headings", height=10)
        placeholder_list.heading("#0", text="Placeholder")
        placeholder_list.heading("original", text="Tên trường gốc")
        placeholder_list.column("#0", width=250)
        placeholder_list.column("original", width=250)

        scrollbar = ttk.Scrollbar(placeholder_frame, orient="vertical", command=placeholder_list.yview)
        scrollbar.pack(side="right", fill="y")
        placeholder_list.configure(yscrollcommand=scrollbar.set)
        placeholder_list.pack(side="left", fill="both", expand=True)

        placeholders = {normalize_vietnamese(field): field for field in self.fields}
        for placeholder, original in sorted(placeholders.items()):
            placeholder_list.insert("", "end", text=f"{{{{ {placeholder} }}}}", values=(original,))

        context_menu = tk.Menu(popup, tearoff=0)
        context_menu.add_command(label="Sao chép Placeholder", command=lambda: self.copy_placeholder_from_popup(placeholder_list))

        def show_context_menu(event):
            item = placeholder_list.identify_row(event.y)
            if item:
                placeholder_list.selection_set(item)
                context_menu.post(event.x_root, event.y_root)

        placeholder_list.bind("<Button-3>", show_context_menu)
        placeholder_list.bind("<Double-1>", lambda event: self.copy_placeholder_from_popup(placeholder_list))
        placeholder_list.bind("<Control-c>", lambda event: self.copy_placeholder_from_popup(placeholder_list))

        button_frame = ttk.Frame(popup)
        button_frame.pack(fill="x", pady=5)
        ttk.Button(button_frame, text="Xuất Placeholder", command=self.export_placeholders, style="primary.TButton").pack(side="left", padx=5)
        ttk.Button(button_frame, text="Đóng", command=popup.destroy, style="danger.TButton").pack(side="right", padx=5)

    def copy_placeholder_from_popup(self, placeholder_list):
        selected = placeholder_list.selection()
        if selected:
            placeholder = placeholder_list.item(selected[0])["text"]
            self.root.clipboard_clear()
            self.root.clipboard_append(placeholder)
            
    def export_placeholders(self):
        """Xuất danh sách placeholder ra file Word hoặc text."""
        if not self.current_config_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn cấu hình trước!")
            return

        placeholders = {}
        for field in self.fields:
            normalized = normalize_vietnamese(field)
            placeholders[normalized] = field

        popup = create_centered_popup(self.root, "Xuất Placeholder", 300, 150) 
        popup.bind('<Escape>', lambda e: popup.destroy())
        ttk.Label(popup, text="Chọn định dạng xuất:").pack(pady=10)
        file_format = tk.StringVar(value="Word")
        ttk.Radiobutton(popup, text="Word (.docx)", variable=file_format, value="Word").pack(anchor="w", padx=20)
        ttk.Radiobutton(popup, text="Text (.txt)", variable=file_format, value="Text").pack(anchor="w", padx=20)

        def confirm_export():
            output_path = filedialog.asksaveasfilename(
                defaultextension=f".{file_format.get().lower()}",
                filetypes=[("Word files", "*.docx") if file_format.get() == "Word" else ("Text files", "*.txt")],
                initialfile=f"placeholders_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{file_format.get().lower()}"
            )
            if not output_path:
                popup.destroy()
                return

            try:
                if file_format.get() == "Word":
                    doc = Document()
                    doc.add_heading("Danh sách Placeholder", level=1)
                    doc.add_paragraph("Sao chép các placeholder dưới đây và dán vào template Word theo cú pháp {{ placeholder }}:")
                    table = doc.add_table(rows=1, cols=2)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "Placeholder"
                    hdr_cells[1].text = "Tên trường gốc"
                    for placeholder, original in sorted(placeholders.items()):
                        row_cells = table.add_row().cells
                        row_cells[0].text = f"{{{{ {placeholder} }}}}"
                        row_cells[1].text = original
                    doc.save(output_path)
                else:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write("Danh sách Placeholder\n")
                        f.write("Sao chép các placeholder dưới đây và dán vào template Word theo cú pháp {{ placeholder }}:\n\n")
                        f.write("Placeholder\tTên trường gốc\n")
                        f.write("-" * 50 + "\n")
                        for placeholder, original in sorted(placeholders.items()):
                            f.write(f"{{{{ {placeholder} }}}}\t{original}\n")
                os.startfile(output_path)
                messagebox.showinfo("Thành công", f"Đã xuất danh sách placeholder ra: {output_path}")
                logging.info(f"Xuất danh sách placeholder: {output_path}")
                popup.destroy()
            except Exception as e:
                logging.error(f"Lỗi khi xuất placeholder: {str(e)}")
                messagebox.showerror("Lỗi", f"Không thể xuất file: {str(e)}")
                popup.destroy()

        ttk.Button(popup, text="Xuất", command=confirm_export, style="primary.TButton").pack(pady=10)

    def check_template_placeholders(self, doc_paths, data_lower):
        try:
            placeholders = set()
            for doc_path in doc_paths:
                doc = DocxTemplate(doc_path)
                for item in doc.get_undeclared_template_variables():
                    # Bỏ qua các placeholder liên quan đến thành viên (dạng ho_ten_1, von_gop_1, v.v.)
                    if any(field in item for field in self.member_columns) and any(item.endswith(f"_{i}") for i in range(1, 100)):
                        continue
                    placeholders.add(item)
            missing_fields = [p for p in placeholders if p not in data_lower]
            return missing_fields
        except Exception as e:
            logging.error(f"Lỗi khi kiểm tra placeholder: {str(e)}")
            messagebox.showerror("Lỗi Template", f"Lỗi khi kiểm tra placeholder:\n{str(e)}")
            return None

#================================ KHU VỰC QUẢN LÝ TEMPLATES ========================================
            
    def update_template_tree(self):
        """Cập nhật Treeview với templates của cấu hình hiện tại."""
        for item in self.template_tree.get_children():
            self.template_tree.delete(item)
        if self.current_config_name and self.current_config_name in self.configs:
            templates = self.configs[self.current_config_name].get("templates", {})
            for template in templates.keys():
                self.template_tree.insert("", "end", text=template)

    def drop_template_files(self, event):
        """Thêm template bằng kéo thả, sao chép file vào thư mục templates và tránh ghi đè."""
        if not self.current_config_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn cấu hình trước!")
            return
        files = self.root.tk.splitlist(event.data)
        added_count = 0
        templates = self.configs[self.current_config_name].get("templates", {})
        for file_path in files:
            if file_path.endswith('.docx') and os.path.exists(file_path):
                template_name = os.path.basename(file_path)
                base_name = os.path.splitext(template_name)[0]
                extension = os.path.splitext(template_name)[1]
                new_name = base_name + extension
                counter = 1
                target_path = os.path.join(self.templates_dir, new_name)
                while os.path.exists(target_path):
                    new_name = f"{base_name}_{counter}{extension}"
                    target_path = os.path.join(self.templates_dir, new_name)
                    counter += 1
                while new_name in templates:
                    new_name = f"{base_name}_{counter}{extension}"
                    target_path = os.path.join(self.templates_dir, new_name)
                    counter += 1
                try:
                    import shutil
                    shutil.copy2(file_path, target_path)
                    templates[new_name] = new_name
                    added_count += 1
                except Exception as e:
                    logging.error(f"Lỗi khi sao chép template {file_path}: {str(e)}")
                    messagebox.showerror("Lỗi", f"Không thể sao chép template: {str(e)}")
        if added_count > 0:
            self.configs[self.current_config_name]["templates"] = templates
            self.save_configs()
            self.update_template_tree()
            messagebox.showinfo("Thành công", f"Đã thêm {added_count} template vào cấu hình '{self.current_config_name}'!")
            logging.info(f"Thêm {added_count} template vào cấu hình '{self.current_config_name}'")

    def add_multiple_templates(self):
        """Thêm nhiều template, sao chép file vào thư mục templates và tránh ghi đè."""
        if not self.current_config_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn cấu hình trước!")
            return
        template_paths = filedialog.askopenfilenames(filetypes=[("Word files", "*.docx")], title="Chọn nhiều template Word")
        if template_paths:
            added_count = 0
            templates = self.configs[self.current_config_name].get("templates", {})
            for template_path in template_paths:
                if os.path.exists(template_path):
                    template_name = os.path.basename(template_path)
                    base_name = os.path.splitext(template_name)[0]
                    extension = os.path.splitext(template_name)[1]
                    new_name = base_name + extension
                    counter = 1
                    target_path = os.path.join(self.templates_dir, new_name)
                    while os.path.exists(target_path):
                        new_name = f"{base_name}_{counter}{extension}"
                        target_path = os.path.join(self.templates_dir, new_name)
                        counter += 1
                    while new_name in templates:
                        new_name = f"{base_name}_{counter}{extension}"
                        target_path = os.path.join(self.templates_dir, new_name)
                        counter += 1
                    try:
                        import shutil
                        shutil.copy2(template_path, target_path)
                        templates[new_name] = new_name
                        added_count += 1
                    except Exception as e:
                        logging.error(f"Lỗi khi sao chép template {template_path}: {str(e)}")
                        messagebox.showerror("Lỗi", f"Không thể sao chép template: {str(e)}")
            if added_count > 0:
                self.configs[self.current_config_name]["templates"] = templates
                self.save_configs()
                self.update_template_tree()
                messagebox.showinfo("Thành công", f"Đã thêm {added_count} template vào cấu hình '{self.current_config_name}'!")
                logging.info(f"Thêm {added_count} template vào cấu hình '{self.current_config_name}'")
   
    def delete_template(self):
        """Xóa template khỏi cấu hình hiện tại."""
        if not self.current_config_name:
            return
        selected_items = self.template_tree.selection()
        if not selected_items:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn template để xóa!")
            return
        if messagebox.askyesno("Xác nhận", "Bạn có chắc muốn xóa các template đã chọn không?"):
            templates = self.configs[self.current_config_name].get("templates", {})
            for item in selected_items:
                template_name = self.template_tree.item(item)["text"]
                del templates[template_name]
            self.configs[self.current_config_name]["templates"] = templates
            self.save_configs()
            self.update_template_tree()
            messagebox.showinfo("Thành công", "Đã xóa các template đã chọn!")
            logging.info(f"Xóa template đã chọn khỏi cấu hình '{self.current_config_name}'")

    def show_template_context_menu(self, event):
        context_menu = tk.Menu(self.root, tearoff=0)
        context_menu.add_command(label="Thêm template", command=self.add_multiple_templates)
        context_menu.add_command(label="Xóa", command=self.delete_template)
        context_menu.post(event.x_root, event.y_root)

    # kéo thả templates
    def start_drag(self, event):
        item = self.template_tree.identify_row(event.y)
        if item:
            self.drag_item = item

    def drag_template(self, event):
        if self.drag_item:
            self.template_tree.selection_set(self.drag_item)

    def drop_template(self, event):
        if self.drag_item:
            target = self.template_tree.identify_row(event.y)
            if target and target != self.drag_item:
                templates = list(self.configs[self.current_config_name].get("templates", {}).keys())
                dragged_name = self.template_tree.item(self.drag_item)["text"]
                target_name = self.template_tree.item(target)["text"]
                dragged_idx = templates.index(dragged_name)
                target_idx = templates.index(target_name)
                templates.insert(target_idx, templates.pop(dragged_idx))
                new_templates = {templates[i]: templates[i] for i in range(len(templates))}
                self.configs[self.current_config_name]["templates"] = new_templates
                self.save_configs()
                self.update_template_tree()
            self.drag_item = None
            
#================================ KHU VỰC QUẢN LÝ NHẬP LIỆU ========================================

    def create_tabs(self):
        """Tạo các tab dựa trên field_groups, bao gồm tab Thông tin thành viên và Ngành nghề kinh doanh."""
        tab_names = list(self.field_groups.keys())
        for i, tab_name in enumerate(tab_names):
            if tab_name == "Ngành nghề kinh doanh":
                self.create_industry_tab()
            elif tab_name == "Thông tin thành viên":
                self.create_member_tab()
            else:
                tab = ttk.Frame(self.notebook)
                self.notebook.add(tab, text=tab_name)
                canvas = tk.Canvas(tab, bg="#ffffff", highlightthickness=0)
                scrollbar = ttk.Scrollbar(tab, orient="vertical", command=canvas.yview)
                scrollable_frame = ttk.Frame(canvas, padding=5)
                scrollable_frame.bind("<Configure>", lambda e, c=canvas: c.configure(scrollregion=c.bbox("all")))
                canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
                canvas.configure(yscrollcommand=scrollbar.set)
                canvas.pack(side="left", fill="both", expand=True)
                scrollbar.pack(side="right", fill="y")
                
                scrollable_frame.grid_columnconfigure(0, weight=1)
                scrollable_frame.grid_columnconfigure(1, weight=2)
                
                # Đảm bảo mỗi tab có đủ trường để kiểm tra cuộn (tạm thời thêm nhiều trường giả)
                fields = self.field_groups[tab_name]
                for j, field in enumerate(fields):
                    display_name = field.replace('_', ' ').title()
                    label_width = max(60, len(display_name) // 2)
                    label = ttk.Label(scrollable_frame, text=f"{display_name}:", width=label_width, anchor="e", wraplength=500)
                    label.grid(row=j, column=0, padx=5, pady=2, sticky="e")
                    self.labels[field] = label
                    entry = ttk.Entry(scrollable_frame, width=50)
                    entry.grid(row=j, column=1, padx=5, pady=2, sticky="ew")
                    self.entries[field] = entry

                    # Gọi phương thức để thêm menu ngữ cảnh cho ô nhập liệu
                    self.add_entry_context_menu(entry)

                    # Lấy menu ngữ cảnh hiện có của entry
                    context_menu = entry.context_menu if hasattr(entry, 'context_menu') else tk.Menu(self.root, tearoff=0)
                    entry.context_menu = context_menu

                    # Thêm các tùy chọn quản lý trường vào menu ngữ cảnh
                    context_menu.add_separator()
                    context_menu.add_command(label="Thêm trường", image=self.add_icon_img, compound="left", command=self.add_field)
                    context_menu.add_command(label="Xóa trường", image=self.delete_icon_img, compound="left", command=lambda f=field: self.delete_field(f))
                    context_menu.add_command(label="Sửa tên trường", image=self.edit_icon_img, compound="left", command=lambda f=field: self.rename_field(f))
                    
                    # Hàm hiển thị menu ngữ cảnh
                    def show_context_menu(event, menu=context_menu):
                        menu.post(event.x_root, event.y_root)

                    # Gán menu ngữ cảnh cho cả label và entry
                    label.bind("<Button-3>", show_context_menu)
                    entry.bind("<Button-3>", show_context_menu)

                    if field == "von_đieu_le":
                        def update_von_dieu_le_bang_chu(event):
                            von_dieu_le_value = self.entries["von_đieu_le"].get()
                            von_dieu_le_bang_chu = number_to_words(von_dieu_le_value)
                            if "von_đieu_le_bang_chu" in self.entries:
                                self.entries["von_đieu_le_bang_chu"].delete(0, tk.END)
                                self.entries["von_đieu_le_bang_chu"].insert(0, von_dieu_le_bang_chu)
                        entry.bind("<KeyRelease>", update_von_dieu_le_bang_chu)

                # Thêm sự kiện cuộn chuột cho từng canvas riêng biệt
                def on_mousewheel(event, c=canvas):  # Truyền canvas cụ thể vào hàm
                    c.yview_scroll(int(-1 * (event.delta / 120)), "units")
                canvas.bind("<Enter>", lambda e, c=canvas: c.bind_all("<MouseWheel>", lambda evt: on_mousewheel(evt, c)))
                canvas.bind("<Leave>", lambda e, c=canvas: c.unbind_all("<MouseWheel>"))

        if self.current_tab_index < len(tab_names):
            self.notebook.select(self.current_tab_index)
            
        self.update_field_dropdown()  # Cập nhật danh sách trường khi tạo tab

    def add_entry_context_menu(self, entry):
        """Thêm menu ngữ cảnh cho ô nhập liệu."""
        context_menu = tk.Menu(self.root, tearoff=0)
        context_menu.add_command(label="Sao chép", command=lambda: entry.event_generate("<<Copy>>"))
        context_menu.add_command(label="Dán", command=lambda: entry.event_generate("<<Paste>>"))
        context_menu.add_command(label="Xóa", command=lambda: entry.delete(0, tk.END))
        # Lưu menu vào thuộc tính context_menu của entry
        entry.context_menu = context_menu

    def copy_text(self, entry):
        text = entry.get()
        self.root.clipboard_clear()
        self.root.clipboard_append(text)

    def cut_text(self, entry):
        text = entry.get()
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        entry.delete(0, tk.END)

    def paste_text(self, entry):
        try:
            text = self.root.clipboard_get()
            entry.delete(0, tk.END)
            entry.insert(0, text)
        except tk.TclError:
            pass

    def add_tab(self):
        tab_name = simpledialog.askstring("Thêm tab", "Nhập tên tab mới:")
        if tab_name and tab_name not in self.field_groups:
            self.field_groups[tab_name] = []
            self.configs[self.current_config_name]["field_groups"] = self.field_groups
            self.save_configs()
            self.clear_tabs()
            self.create_tabs()
            self.tab_dropdown["values"] = list(self.field_groups.keys())
            self.tab_dropdown.set(tab_name)
            messagebox.showinfo("Thành công", f"Đã thêm tab '{tab_name}'!")
            logging.info(f"Thêm tab '{tab_name}'")

    def delete_tab(self):
        selected_tab = self.tab_var.get()
        if not selected_tab:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn tab để xóa!")
            return
        if messagebox.askyesno("Xác nhận", f"Bạn có chắc muốn xóa tab '{selected_tab}' không?"):
            for field in self.field_groups[selected_tab]:
                self.fields.remove(field)
            del self.field_groups[selected_tab]
            self.configs[self.current_config_name]["field_groups"] = self.field_groups
            self.save_configs()
            self.clear_tabs()
            self.create_tabs()
            self.tab_dropdown["values"] = list(self.field_groups.keys())
            self.tab_dropdown.set(list(self.field_groups.keys())[0] if self.field_groups else "")
            messagebox.showinfo("Thành công", f"Đã xóa tab '{selected_tab}'!")
            logging.info(f"Xóa tab '{selected_tab}'")

    def rename_tab(self):
        old_tab = self.tab_var.get()
        if not old_tab:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn tab để sửa tên!")
            return
        new_tab = simpledialog.askstring("Sửa tên tab", "Nhập tên mới:", initialvalue=old_tab)
        if new_tab and new_tab != old_tab and new_tab not in self.field_groups:
            self.field_groups[new_tab] = self.field_groups.pop(old_tab)
            self.configs[self.current_config_name]["field_groups"] = self.field_groups
            self.save_configs()
            self.clear_tabs()
            self.create_tabs()
            self.tab_dropdown["values"] = list(self.field_groups.keys())
            self.tab_dropdown.set(new_tab)
            messagebox.showinfo("Thành công", f"Đã đổi tên tab thành '{new_tab}'!")
            logging.info(f"Đổi tên tab từ '{old_tab}' thành '{new_tab}'")

    def clear_entries(self):
        for entry in self.entries.values():
            entry.delete(0, tk.END)

    def clear_tabs(self): 
        """Xóa tất cả các tab trong notebook."""
        current_tab_index = self.notebook.index("current") if self.notebook.tabs() else 0
        for tab in self.notebook.tabs():
            self.notebook.forget(tab)
        self.notebook.update_idletasks()
        self.entries.clear()
        self.labels.clear()
        self.current_tab_index = current_tab_index

    def create_member_tab(self):
        """Tạo tab Thông tin thành viên với Treeview và scrollbar dọc."""
        tab_name = "Thông tin thành viên"
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text=tab_name)
        
        # Tạo Canvas và Scrollbar để hỗ trợ cuộn dọc
        canvas = tk.Canvas(tab, bg="#ffffff", highlightthickness=0)
        scrollbar_y = ttk.Scrollbar(tab, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, padding=5)
        scrollable_frame.bind("<Configure>", lambda e, c=canvas: c.configure(scrollregion=c.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar_y.set)
        
        # Đặt vị trí Canvas và Scrollbar dọc
        canvas.pack(side="left", fill="both", expand=True, padx=10, pady=10)
        scrollbar_y.pack(side="right", fill="y")
        
        # Tạo frame chứa Treeview
        tree_frame = ttk.Frame(scrollable_frame)
        tree_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Chỉ hiển thị các cột quan trọng
        self.displayed_columns = ["ho_ten", "so_cccd", "dia_chi_thuong_tru", "von_gop", "ty_le_gop", "la_chu_tich"]
        self.member_tree = ttk.Treeview(tree_frame, columns=self.displayed_columns, show="headings", height=15)  # Tăng height để hiển thị nhiều hàng hơn
        
        # Định nghĩa tiêu đề và chiều rộng cột
        column_widths = {
            "ho_ten": 200,      # Họ tên rộng hơn để dễ đọc
            "so_cccd": 100,
            "dia_chi_thuong_tru": 400,
            "von_gop": 120,
            "ty_le_gop": 100,
            "la_chu_tich": 80,  # Cột Chủ tịch nhỏ gọn
        }
        for col in self.displayed_columns:
            self.member_tree.heading(col, text=col.replace('_', ' ').title())
            width = column_widths.get(col, 100)
            self.member_tree.column(col, width=width, anchor="center")
        
        # Đặt Treeview vào frame
        self.member_tree.pack(side="top", fill="both", expand=True)
        
        # Tạo menu ngữ cảnh (context menu)
        context_menu = tk.Menu(self.root, tearoff=0)
        context_menu.add_command(label="Thêm thành viên", command=self.add_member)
        context_menu.add_command(label="Sửa thành viên", command=self.edit_member)
        context_menu.add_command(label="Xóa thành viên", command=self.delete_member)
        context_menu.add_command(label="Xem chi tiết", command=self.view_member_details)

        def show_context_menu(event):
            selected = self.member_tree.identify_row(event.y)
            if selected:
                self.member_tree.selection_set(selected)
            context_menu.post(event.x_root, event.y_root)

        # Gắn sự kiện nhấp chuột phải để hiển thị menu ngữ cảnh
        self.member_tree.bind("<Button-3>", show_context_menu)
        
        # Gắn sự kiện nhấp đúp để xem chi tiết
        self.member_tree.bind("<Double-1>", self.view_member_details)

        # Gắn các sự kiện kéo thả
        self.member_tree.bind("<Button-1>", self.start_drag_member)
        self.member_tree.bind("<B1-Motion>", self.drag_member)
        self.member_tree.bind("<ButtonRelease-1>", self.drop_member)
        
        # Tạo frame cho các nút
        button_frame = ttk.Frame(scrollable_frame)
        button_frame.pack(pady=10)
        
        ttk.Button(button_frame, text="Thêm thành viên", command=self.add_member).pack(side="left", padx=10, expand=True)
        ttk.Button(button_frame, text="Sửa thành viên", command=self.edit_member).pack(side="left", padx=10, expand=True)
        ttk.Button(button_frame, text="Xóa thành viên", command=self.delete_member).pack(side="left", padx=10, expand=True)
        ttk.Button(button_frame, text="Xem chi tiết", command=self.view_member_details).pack(side="left", padx=10, expand=True)
        
        # Kích hoạt cuộn chuột dọc
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        canvas.bind("<Enter>", lambda e: canvas.bind_all("<MouseWheel>", on_mousewheel))
        canvas.bind("<Leave>", lambda e: canvas.unbind_all("<MouseWheel>"))
        
        # Gọi load_member_data để hiển thị dữ liệu ngay khi tab được tạo
        self.load_member_data()

    def load_member_data(self):
        """Tải dữ liệu thành viên vào Treeview."""
        self.member_tree.delete(*self.member_tree.get_children())
        selected_name = self.load_data_var.get()
        members = []
        for entry in self.saved_entries:
            if entry["name"] == selected_name:
                members = entry["data"].get("thanh_vien", [])
                break
        for member in members:
            values = [member.get(col, "") for col in self.displayed_columns]
            values[self.displayed_columns.index("la_chu_tich")] = "X" if member.get("la_chu_tich", False) else ""
            self.member_tree.insert("", "end", values=values)
       
    def add_member(self):
        """Thêm một thành viên mới với tab và checkbox chức danh chủ tịch."""
        popup = create_centered_popup(self.root, "Thêm thành viên", 450, 550)
        popup.bind('<Escape>', lambda e: popup.destroy())

        # Tạo notebook để chia tab
        notebook = ttk.Notebook(popup)
        notebook.pack(fill="both", expand=True, padx=5, pady=5)

        # Tab 1: Thông tin cá nhân
        tab1 = ttk.Frame(notebook)
        notebook.add(tab1, text="Thông tin cá nhân")
        tab1_fields = ["ho_ten", "gioi_tinh", "ngay_sinh", "dan_toc", "quoc_tich"]
        entries = {}
        for col in tab1_fields:
            ttk.Label(tab1, text=col.replace('_', ' ').title()).pack(pady=5)
            entry = ttk.Entry(tab1, width=50)
            entry.pack(pady=5)
            entries[col] = entry

        # Tab 2: Thông tin giấy tờ
        tab2 = ttk.Frame(notebook)
        notebook.add(tab2, text="Thông tin giấy tờ")
        tab2_fields = ["loai_giay_to", "so_cccd", "ngay_cap", "noi_cap", "ngay_het_han"]
        for col in tab2_fields:
            ttk.Label(tab2, text=col.replace('_', ' ').title()).pack(pady=5)
            entry = ttk.Entry(tab2, width=50)
            entry.pack(pady=5)
            entries[col] = entry

        # Tab 3: Thông tin địa chỉ
        tab3 = ttk.Frame(notebook)
        notebook.add(tab3, text="Thông tin địa chỉ")
        tab3_fields = ["dia_chi_thuong_tru", "dia_chi_lien_lac"]
        for col in tab3_fields:
            ttk.Label(tab3, text=col.replace('_', ' ').title()).pack(pady=5)
            entry = ttk.Entry(tab3, width=50)
            entry.pack(pady=5)
            entries[col] = entry

        # Tab 4: Thông tin vốn góp
        tab4 = ttk.Frame(notebook)
        notebook.add(tab4, text="Thông tin vốn góp")
        tab4_fields = ["von_gop", "ty_le_gop", "ngay_gop_von"]
        for col in tab4_fields:
            ttk.Label(tab4, text=col.replace('_', ' ').title()).pack(pady=5)
            entry = ttk.Entry(tab4, width=50)
            entry.pack(pady=5)
            entries[col] = entry

        # Thêm checkbox "Chức danh chủ tịch"
        chairman_var = tk.BooleanVar(value=False)
        chairman_check = ttk.Checkbutton(popup, text="Chức danh chủ tịch", variable=chairman_var)
        chairman_check.pack(pady=5)

        # Nút xác nhận
        def confirm_add():
            member = {col: entries[col].get() for col in self.member_columns}
            member["la_chu_tich"] = chairman_var.get()
            selected_name = self.load_data_var.get()
            for entry in self.saved_entries:
                if entry["name"] == selected_name:
                    members = entry["data"].setdefault("thanh_vien", [])
                    # Nếu thành viên mới là chủ tịch, bỏ chọn chủ tịch cũ
                    if member["la_chu_tich"]:
                        for existing_member in members:
                            if existing_member.get("la_chu_tich", False):
                                existing_member["la_chu_tich"] = False
                        members.insert(0, member)  # Đặt chủ tịch lên đầu
                    else:
                        members.append(member)
                    self.save_configs()
                    self.load_member_data()
                    popup.destroy()
                    break

        ttk.Button(popup, text="Thêm", command=confirm_add).pack(pady=10)

    def delete_member(self):
        """Xóa nhiều thành viên đã chọn."""
        selected_items = self.member_tree.selection()
        if not selected_items:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn ít nhất một thành viên để xóa!")
            return

        if messagebox.askyesno("Xác nhận", f"Bạn có muốn xóa {len(selected_items)} thành viên đã chọn không?"):
            selected_name = self.load_data_var.get()
            if not selected_name:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn một mục dữ liệu trước!")
                return

            for entry in self.saved_entries:
                if entry["name"] == selected_name:
                    if "thanh_vien" not in entry["data"] or not entry["data"]["thanh_vien"]:
                        messagebox.showwarning("Cảnh báo", "Không có thành viên nào để xóa!")
                        return

                    indices = [self.member_tree.index(item) for item in selected_items]
                    indices.sort(reverse=True)

                    for idx in indices:
                        try:
                            entry["data"]["thanh_vien"].pop(idx)
                        except IndexError:
                            continue

                    self.save_configs()
                    self.load_member_data()
                    break
            else:
                messagebox.showwarning("Cảnh báo", "Không tìm thấy mục dữ liệu tương ứng!")

    def edit_member(self):
        """Sửa thông tin thành viên đã chọn với tab."""
        selected_item = self.member_tree.selection()
        if not selected_item:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn thành viên để sửa!")
            return
        idx = self.member_tree.index(selected_item)
        selected_name = self.load_data_var.get()
        member = None
        for entry in self.saved_entries:
            if entry["name"] == selected_name:
                member = entry["data"]["thanh_vien"][idx]
                break
        
        popup = create_centered_popup(self.root, "Sửa thành viên", 450, 550)
        popup.bind('<Escape>', lambda e: popup.destroy())

        # Tạo notebook để chia tab
        notebook = ttk.Notebook(popup)
        notebook.pack(fill="both", expand=True, padx=5, pady=5)

        # Tab 1: Thông tin cá nhân
        tab1 = ttk.Frame(notebook)
        notebook.add(tab1, text="Thông tin cá nhân")
        tab1_fields = ["ho_ten", "gioi_tinh", "ngay_sinh", "dan_toc", "quoc_tich"]
        entries = {}
        for col in tab1_fields:
            ttk.Label(tab1, text=col.replace('_', ' ').title()).pack(pady=5)
            entry = ttk.Entry(tab1, width=50)
            entry.insert(0, member.get(col, ""))
            entry.pack(pady=5)
            entries[col] = entry

        # Tab 2: Thông tin giấy tờ
        tab2 = ttk.Frame(notebook)
        notebook.add(tab2, text="Thông tin giấy tờ")
        tab2_fields = ["loai_giay_to", "so_cccd", "ngay_cap", "noi_cap", "ngay_het_han"]
        for col in tab2_fields:
            ttk.Label(tab2, text=col.replace('_', ' ').title()).pack(pady=5)
            entry = ttk.Entry(tab2, width=50)
            entry.insert(0, member.get(col, ""))
            entry.pack(pady=5)
            entries[col] = entry

        # Tab 3: Thông tin địa chỉ
        tab3 = ttk.Frame(notebook)
        notebook.add(tab3, text="Thông tin địa chỉ")
        tab3_fields = ["dia_chi_thuong_tru", "dia_chi_lien_lac"]
        for col in tab3_fields:
            ttk.Label(tab3, text=col.replace('_', ' ').title()).pack(pady=5)
            entry = ttk.Entry(tab3, width=50)
            entry.insert(0, member.get(col, ""))
            entry.pack(pady=5)
            entries[col] = entry

        # Tab 4: Thông tin vốn góp
        tab4 = ttk.Frame(notebook)
        notebook.add(tab4, text="Thông tin vốn góp")
        tab4_fields = ["von_gop", "ty_le_gop", "ngay_gop_von"]
        for col in tab4_fields:
            ttk.Label(tab4, text=col.replace('_', ' ').title()).pack(pady=5)
            entry = ttk.Entry(tab4, width=50)
            entry.insert(0, member.get(col, ""))
            entry.pack(pady=5)
            entries[col] = entry

        # Thêm checkbox "Chức danh chủ tịch"
        chairman_var = tk.BooleanVar(value=member.get("la_chu_tich", False))
        chairman_check = ttk.Checkbutton(popup, text="Chức danh chủ tịch", variable=chairman_var)
        chairman_check.pack(pady=5)

        # Nút xác nhận
        def confirm_edit():
            updated_member = {col: entries[col].get() for col in self.member_columns}
            updated_member["la_chu_tich"] = chairman_var.get()
            for entry in self.saved_entries:
                if entry["name"] == selected_name:
                    members = entry["data"]["thanh_vien"]
                    # Nếu thành viên được sửa thành chủ tịch, bỏ chọn chủ tịch cũ
                    if updated_member["la_chu_tich"] and not member.get("la_chu_tich", False):
                        for i, existing_member in enumerate(members):
                            if i != idx and existing_member.get("la_chu_tich", False):
                                existing_member["la_chu_tich"] = False
                        # Đưa thành viên lên đầu danh sách nếu là chủ tịch
                        members.pop(idx)
                        members.insert(0, updated_member)
                    else:
                        members[idx] = updated_member
                    self.save_configs()
                    self.load_member_data()
                    popup.destroy()
                    break
        
        ttk.Button(popup, text="Lưu", command=confirm_edit).pack(pady=10)

    def view_member_details(self, event=None):
        """Hiển thị chi tiết thông tin thành viên khi nhấp đúp hoặc chọn từ menu ngữ cảnh."""
        selected_item = self.member_tree.selection()
        if not selected_item:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn thành viên để xem chi tiết!")
            return
        idx = self.member_tree.index(selected_item)
        selected_name = self.load_data_var.get()
        member = None
        for entry in self.saved_entries:
            if entry["name"] == selected_name:
                member = entry["data"]["thanh_vien"][idx]
                break
        
        # Tạo popup hiển thị chi tiết
        popup = create_centered_popup(self.root, "Chi tiết thành viên", 500, 600)
        popup.bind('<Escape>', lambda e: popup.destroy())

        # Tạo notebook để chia tab
        notebook = ttk.Notebook(popup)
        notebook.pack(fill="both", expand=True, padx=5, pady=5)

        # Tab 1: Thông tin cá nhân
        tab1 = ttk.Frame(notebook)
        notebook.add(tab1, text="Thông tin cá nhân")
        tab1_fields = ["ho_ten", "gioi_tinh", "ngay_sinh", "dan_toc", "quoc_tich"]
        for col in tab1_fields:
            ttk.Label(tab1, text=col.replace('_', ' ').title() + ":").pack(pady=5)
            ttk.Label(tab1, text=member.get(col, ""), wraplength=450).pack(pady=5)

        # Tab 2: Thông tin giấy tờ
        tab2 = ttk.Frame(notebook)
        notebook.add(tab2, text="Thông tin giấy tờ")
        tab2_fields = ["loai_giay_to", "so_cccd", "ngay_cap", "noi_cap", "ngay_het_han"]
        for col in tab2_fields:
            ttk.Label(tab2, text=col.replace('_', ' ').title() + ":").pack(pady=5)
            ttk.Label(tab2, text=member.get(col, ""), wraplength=450).pack(pady=5)

        # Tab 3: Thông tin địa chỉ
        tab3 = ttk.Frame(notebook)
        notebook.add(tab3, text="Thông tin địa chỉ")
        tab3_fields = ["dia_chi_thuong_tru", "dia_chi_lien_lac"]
        for col in tab3_fields:
            ttk.Label(tab3, text=col.replace('_', ' ').title() + ":").pack(pady=5)
            ttk.Label(tab3, text=member.get(col, ""), wraplength=450).pack(pady=5)

        # Tab 4: Thông tin vốn góp
        tab4 = ttk.Frame(notebook)
        notebook.add(tab4, text="Thông tin vốn góp")
        tab4_fields = ["von_gop", "ty_le_gop", "ngay_gop_von"]
        for col in tab4_fields:
            ttk.Label(tab4, text=col.replace('_', ' ').title() + ":").pack(pady=5)
            ttk.Label(tab4, text=member.get(col, ""), wraplength=450).pack(pady=5)

        # Hiển thị trạng thái chủ tịch
        ttk.Label(popup, text="Chức danh: " + ("Chủ tịch" if member.get("la_chu_tich", False) else "Thành viên")).pack(pady=5)
        def open_edit_and_close():
            popup.destroy()  # Đóng popup hiện tại
            self.edit_member()
        # Nút đóng
        ttk.Button(popup, text="Đóng", command=popup.destroy).pack(side="right", pady=10, expand=True)
        ttk.Button(popup, text="Chỉnh sửa", command=open_edit_and_close).pack(side="left", pady=10, expand=True)

    def start_drag_member(self, event):
        """Bắt đầu kéo một thành viên."""
        item = self.member_tree.identify_row(event.y)
        if item:
            self.drag_item = item

    def drag_member(self, event):
        """Di chuyển thành viên khi kéo."""
        if self.drag_item:
            self.member_tree.selection_set(self.drag_item)

    def drop_member(self, event):
        """Thả thành viên vào vị trí mới và cập nhật danh sách."""
        if self.drag_item:
            target = self.member_tree.identify_row(event.y)
            if target and target != self.drag_item:
                selected_name = self.load_data_var.get()
                for entry in self.saved_entries:
                    if entry["name"] == selected_name:
                        members = entry["data"].get("thanh_vien", [])
                        dragged_idx = self.member_tree.index(self.drag_item)
                        target_idx = self.member_tree.index(target)
                        dragged_member = members[dragged_idx]
                        # Nếu thành viên được kéo là chủ tịch, giữ nó ở đầu
                        if dragged_member.get("la_chu_tich", False):
                            messagebox.showwarning("Cảnh báo", "Chủ tịch phải ở vị trí đầu tiên!")
                            break
                        # Di chuyển thành viên trong danh sách
                        members.insert(target_idx, members.pop(dragged_idx))
                        self.save_configs()
                        self.load_member_data()
                        break
            self.drag_item = None
            
    def create_industry_tab(self):
        """Tạo tab Ngành nghề kinh doanh với Treeview và các nút quản lý."""
        tab_name = "Ngành nghề kinh doanh"
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text=tab_name)
        
        # Tạo Canvas và Scrollbar để hỗ trợ cuộn
        canvas = tk.Canvas(tab, bg="#ffffff", highlightthickness=0)
        scrollbar = ttk.Scrollbar(tab, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas, padding=5)
        scrollable_frame.bind("<Configure>", lambda e, c=canvas: c.configure(scrollregion=c.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="center")  # Căn giữa scrollable_frame
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Căn giữa canvas và scrollbar trong tab
        canvas.pack(side="left", fill="both", expand=True, padx=20, pady=20)  # Thêm padding để tạo khoảng cách
        scrollbar.pack(side="right", fill="y")
        
        # Tăng kích thước Treeview và căn giữa
        self.industry_tree = ttk.Treeview(scrollable_frame, columns=("ten_nganh", "ma_nganh", "la_nganh_chinh"), show="headings", height=15)
        self.industry_tree.heading("ten_nganh", text="Tên ngành")
        self.industry_tree.heading("ma_nganh", text="Mã ngành")
        self.industry_tree.heading("la_nganh_chinh", text="Ngành chính")
        
        # Tăng chiều rộng cột để hiển thị nội dung rõ ràng hơn
        self.industry_tree.column("ten_nganh", width=600)  # Tăng từ 400 lên 600
        self.industry_tree.column("ma_nganh", width=150)   # Tăng từ 100 lên 150
        self.industry_tree.column("la_nganh_chinh", width=150)  # Tăng từ 100 lên 150
        
        # Căn giữa Treeview trong scrollable_frame
        self.industry_tree.pack(fill="x", expand=True, padx=10, pady=10)
        
        # Tạo menu ngữ cảnh (context menu)
        context_menu = tk.Menu(self.root, tearoff=0)
        context_menu.add_command(label="Thêm ngành", command=self.add_industry)
        context_menu.add_command(label="Sửa ngành", command=self.edit_industry)
        context_menu.add_command(label="Xóa ngành", command=self.delete_industry)
        context_menu.add_command(label="Ngành chính", command=self.set_main_industry)  # Thêm tùy chọn "Ngành chính"
        context_menu.add_command(label="Xem chi tiết", command=self.view_industry_details)

        def show_context_menu(event):
            selected = self.industry_tree.identify_row(event.y)
            if selected:
                self.industry_tree.selection_set(selected)
            context_menu.post(event.x_root, event.y_root)

        # Gắn sự kiện nhấp chuột phải để hiển thị menu ngữ cảnh
        self.industry_tree.bind("<Button-3>", show_context_menu)
        
        # Gắn sự kiện nhấp đúp để xem chi tiết
        self.industry_tree.bind("<Double-1>", self.view_industry_details)
        
        # Tạo frame cho các nút và căn giữa
        button_frame = ttk.Frame(scrollable_frame)
        button_frame.pack(pady=10)
        
        # Tạo các nút và đặt chúng trong button_frame (xóa nút "Đặt ngành chính")
        ttk.Button(button_frame, text="Thêm ngành", command=self.add_industry).pack(side="left", padx=10, expand=True)
        ttk.Button(button_frame, text="Xóa ngành", command=self.delete_industry).pack(side="left", padx=10, expand=True)
        ttk.Button(button_frame, text="Sửa ngành", command=self.edit_industry).pack(side="left", padx=10, expand=True)
        ttk.Button(button_frame, text="Xem chi tiết", command=self.view_industry_details).pack(side="left", padx=10, expand=True)

        # Kích hoạt cuộn chuột
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        canvas.bind("<Enter>", lambda e: canvas.bind_all("<MouseWheel>", on_mousewheel))
        canvas.bind("<Leave>", lambda e: canvas.unbind_all("<MouseWheel>"))
        
        # Load dữ liệu ngành nghề
        self.load_industry_data()

    def load_industry_data(self):
        """Tải dữ liệu ngành nghề vào Treeview."""
        self.industry_tree.delete(*self.industry_tree.get_children())
        selected_name = self.load_data_var.get()
        industries = []
        for entry in self.saved_entries:
            if entry["name"] == selected_name:
                industries = entry["data"].get("nganh_nghe", [])
                break
        for industry in industries:
            self.industry_tree.insert("", "end", values=(industry["ten_nganh"], industry["ma_nganh"], "X" if industry["la_nganh_chinh"] else ""))

    def add_industry(self):
        """Thêm một ngành nghề mới với tính năng lọc autocomplete và checkbox Ngành chính."""
        # Đường dẫn đến file industry_codes.json trong thư mục AppData
        industry_codes_path = os.path.join(self.appdata_dir, "industry_codes.json")
        
        # Tải danh sách mã ngành từ file
        try:
            with open(industry_codes_path, "r", encoding="utf-8") as f:
                industry_codes = json.load(f)
        except FileNotFoundError:
            messagebox.showerror("Lỗi", f"Không tìm thấy file {industry_codes_path}! Vui lòng tạo file này chứa danh sách mã ngành trong thư mục AppData.")
            return

        popup = create_centered_popup(self.root, "Thêm ngành nghề", 600, 250)
        popup.title("Thêm ngành nghề")
        popup.bind('<Escape>', lambda e: popup.destroy())
        ttk.Label(popup, text="Chọn mã ngành và tên ngành:").pack(pady=5)
        industry_var = tk.StringVar()
        industry_combo = ttk.Combobox(popup, textvariable=industry_var, width=90)
        industry_combo.pack(pady=5)
        
        # Danh sách đầy đủ các mã ngành và tên ngành
        full_list = [f"{code['ma_nganh']} - {code['ten_nganh']}" for code in industry_codes]
        
        # Hàm cập nhật danh sách lọc
        def update_list(*args):
            search_term = industry_var.get().lower()  # Lấy chuỗi tìm kiếm và chuyển thành chữ thường
            if search_term:
                # Lọc các mục chứa chuỗi tìm kiếm (trong mã ngành hoặc tên ngành)
                filtered_list = [item for item in full_list if search_term in item.lower()]
            else:
                # Nếu không có chuỗi tìm kiếm, hiển thị toàn bộ danh sách
                filtered_list = full_list
            industry_combo['values'] = filtered_list  # Cập nhật danh sách trong combobox
        
        # Gắn sự kiện <KeyRelease> để gọi hàm update_list khi người dùng gõ
        industry_combo.bind('<KeyRelease>', update_list)
        
        # Khởi tạo combobox với danh sách đầy đủ ban đầu
        industry_combo['values'] = full_list
        
        ttk.Label(popup, text="Chi tiết ngành (nếu có):").pack(pady=5)
        chi_tiet_var = tk.StringVar()
        chi_tiet_entry = ttk.Entry(popup, textvariable=chi_tiet_var, width=90)
        chi_tiet_entry.pack(pady=5)

        # Thêm checkbox "Ngành chính"
        main_industry_var = tk.BooleanVar(value=False)
        main_industry_check = ttk.Checkbutton(popup, text="Ngành chính", variable=main_industry_var)
        main_industry_check.pack(pady=5)
        
        def confirm_add():
            selected_industry = industry_var.get()
            chi_tiet = chi_tiet_var.get().strip()
            if not selected_industry:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn mã ngành và tên ngành!")
                return
            ma_nganh, ten_nganh = selected_industry.split(" - ", 1)
            if chi_tiet:
                ten_nganh = f"{ten_nganh} - {chi_tiet}"
            selected_name = self.load_data_var.get()
            for entry in self.saved_entries:
                if entry["name"] == selected_name:
                    industries = entry["data"].setdefault("nganh_nghe", [])
                    # Nếu ngành mới là ngành chính, bỏ chọn ngành chính cũ
                    if main_industry_var.get():
                        for existing_industry in industries:
                            if existing_industry.get("la_nganh_chinh", False):
                                existing_industry["la_nganh_chinh"] = False
                    industry = {"ma_nganh": ma_nganh, "ten_nganh": ten_nganh, "la_nganh_chinh": main_industry_var.get()}
                    industries.append(industry)
                    self.save_configs()
                    self.load_industry_data()
                    popup.destroy()
                    break
        
        ttk.Button(popup, text="Thêm", command=confirm_add).pack(pady=5)

    def delete_industry(self):
        """Xóa nhiều ngành nghề đã chọn."""
        selected_items = self.industry_tree.selection()
        if not selected_items:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn ít nhất một ngành nghề để xóa!")
            return

        if messagebox.askyesno("Xác nhận", f"Bạn có muốn xóa {len(selected_items)} ngành nghề đã chọn không?"):
            selected_name = self.load_data_var.get()
            if not selected_name:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn một mục dữ liệu trước!")
                return

            for entry in self.saved_entries:
                if entry["name"] == selected_name:
                    if "nganh_nghe" not in entry["data"] or not entry["data"]["nganh_nghe"]:
                        messagebox.showwarning("Cảnh báo", "Không có ngành nghề nào để xóa!")
                        return

                    indices = [self.industry_tree.index(item) for item in selected_items]
                    indices.sort(reverse=True)

                    for idx in indices:
                        try:
                            entry["data"]["nganh_nghe"].pop(idx)
                        except IndexError:
                            continue

                    self.save_configs()
                    self.load_industry_data()
                    break
            else:
                messagebox.showwarning("Cảnh báo", "Không tìm thấy mục dữ liệu tương ứng!")

    def edit_industry(self, event=None):
        """Sửa chi tiết ngành nghề với bố cục tương tự Thêm ngành và checkbox Ngành chính."""
        # Đường dẫn đến file industry_codes.json trong thư mục AppData
        industry_codes_path = os.path.join(self.appdata_dir, "industry_codes.json")
        
        # Tải danh sách mã ngành từ file
        try:
            with open(industry_codes_path, "r", encoding="utf-8") as f:
                industry_codes = json.load(f)
        except FileNotFoundError:
            messagebox.showerror("Lỗi", f"Không tìm thấy file {industry_codes_path}! Vui lòng tạo file này chứa danh sách mã ngành trong thư mục AppData.")
            return

        selected_item = self.industry_tree.selection()
        if not selected_item:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn ngành nghề để sửa!")
            return
        idx = self.industry_tree.index(selected_item)
        selected_name = self.load_data_var.get()
        industry = None
        for entry in self.saved_entries:
            if entry["name"] == selected_name:
                industry = entry["data"]["nganh_nghe"][idx]
                break
        
        popup = create_centered_popup(self.root, "Sửa ngành nghề", 600, 250)
        popup.bind('<Escape>', lambda e: popup.destroy())
        ttk.Label(popup, text="Chọn mã ngành và tên ngành:").pack(pady=10)
        industry_var = tk.StringVar()
        # Tách chi tiết (nếu có) khỏi tên ngành
        ten_nganh = industry["ten_nganh"]
        chi_tiet = ""
        if " - " in ten_nganh:
            base_ten_nganh, chi_tiet = ten_nganh.rsplit(" - ", 1)
            for code in industry_codes:
                if code["ten_nganh"] == base_ten_nganh:
                    industry_var.set(f"{industry['ma_nganh']} - {base_ten_nganh}")
                    break
        else:
            industry_var.set(f"{industry['ma_nganh']} - {ten_nganh}")
        
        industry_combo = ttk.Combobox(popup, textvariable=industry_var, width=90)
        industry_combo.pack(pady=10)
        
        # Danh sách đầy đủ các mã ngành và tên ngành
        full_list = [f"{code['ma_nganh']} - {code['ten_nganh']}" for code in industry_codes]
        
        # Hàm cập nhật danh sách lọc
        def update_list(*args):
            search_term = industry_var.get().lower()  # Lấy chuỗi tìm kiếm và chuyển thành chữ thường
            if search_term:
                # Lọc các mục chứa chuỗi tìm kiếm (trong mã ngành hoặc tên ngành)
                filtered_list = [item for item in full_list if search_term in item.lower()]
            else:
                # Nếu không có chuỗi tìm kiếm, hiển thị toàn bộ danh sách
                filtered_list = full_list
            industry_combo['values'] = filtered_list  # Cập nhật danh sách trong combobox
        
        # Gắn sự kiện <KeyRelease> để gọi hàm update_list khi người dùng gõ
        industry_combo.bind('<KeyRelease>', update_list)
        
        # Khởi tạo combobox với danh sách đầy đủ ban đầu
        industry_combo['values'] = full_list
        
        ttk.Label(popup, text="Chi tiết ngành (nếu có):").pack(pady=5)
        chi_tiet_var = tk.StringVar(value=chi_tiet)
        chi_tiet_entry = ttk.Entry(popup, textvariable=chi_tiet_var, width=90)
        chi_tiet_entry.pack(pady=10)

        # Thêm checkbox "Ngành chính"
        main_industry_var = tk.BooleanVar(value=industry.get("la_nganh_chinh", False))
        main_industry_check = ttk.Checkbutton(popup, text="Ngành chính", variable=main_industry_var)
        main_industry_check.pack(pady=5)
        
        def confirm_edit():
            selected_industry = industry_var.get()
            chi_tiet = chi_tiet_var.get().strip()
            if not selected_industry:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn mã ngành và tên ngành!")
                return
            ma_nganh, ten_nganh = selected_industry.split(" - ", 1)
            if chi_tiet:
                ten_nganh = f"{ten_nganh} - {chi_tiet}"
            for entry in self.saved_entries:
                if entry["name"] == selected_name:
                    industries = entry["data"]["nganh_nghe"]
                    # Nếu ngành được sửa thành ngành chính, bỏ chọn ngành chính cũ
                    if main_industry_var.get() and not industry.get("la_nganh_chinh", False):
                        for i, existing_industry in enumerate(industries):
                            if i != idx and existing_industry.get("la_nganh_chinh", False):
                                existing_industry["la_nganh_chinh"] = False
                    industries[idx] = {"ma_nganh": ma_nganh, "ten_nganh": ten_nganh, "la_nganh_chinh": main_industry_var.get()}
                    self.save_configs()
                    self.load_industry_data()
                    popup.destroy()
                    break
        
        ttk.Button(popup, text="Lưu", command=confirm_edit).pack(pady=5)

    def view_industry_details(self, event=None):
        """Hiển thị chi tiết thông tin ngành nghề khi nhấp đúp hoặc chọn từ menu ngữ cảnh."""
        selected_item = self.industry_tree.selection()
        if not selected_item:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn ngành nghề để xem chi tiết!")
            return
        idx = self.industry_tree.index(selected_item)
        selected_name = self.load_data_var.get()
        industry = None
        for entry in self.saved_entries:
            if entry["name"] == selected_name:
                industry = entry["data"]["nganh_nghe"][idx]
                break
        
        # Tạo popup chi tiết ngành nghề
        popup = create_centered_popup(self.root, "Chi tiết ngành nghề", 500, 250)
        popup.bind('<Escape>', lambda e: popup.destroy())
         # Hiển thị thông tin ngành nghề
        ttk.Label(popup, text="Tên ngành:").pack(pady=5)
        ttk.Label(popup, text=industry.get("ten_nganh", ""), wraplength=450).pack(pady=5)
        
        ttk.Label(popup, text="Mã ngành:").pack(pady=5)
        ttk.Label(popup, text=industry.get("ma_nganh", "")).pack(pady=5)
        
        ttk.Label(popup, text="Ngành chính:").pack(pady=5)
        ttk.Label(popup, text="Có" if industry.get("la_nganh_chinh", False) else "Không").pack(pady=5)

        # Hàm trung gian để đóng popup trước khi mở edit_industry
        def open_edit_and_close():
            popup.destroy()  # Đóng popup hiện tại
            self.edit_industry()  # Mở popup chỉnh sửa

        # Nút đóng
        ttk.Button(popup, text="Đóng", command=popup.destroy).pack(side="right", padx=5, expand=True)
        ttk.Button(popup, text="Chỉnh sửa", command=open_edit_and_close).pack(side="left", padx=5, expand=True)
        
        
    def set_main_industry(self):
        """Đặt ngành nghề được chọn làm ngành chính."""
        selected_item = self.industry_tree.selection()
        if not selected_item:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn ngành nghề để đặt làm ngành chính!")
            return
        idx = self.industry_tree.index(selected_item)
        selected_name = self.load_data_var.get()
        for entry in self.saved_entries:
            if entry["name"] == selected_name:
                industries = entry["data"].get("nganh_nghe", [])
                # Bỏ chọn ngành chính cũ
                for i, industry in enumerate(industries):
                    if i != idx and industry.get("la_nganh_chinh", False):
                        industry["la_nganh_chinh"] = False
                # Đặt ngành được chọn làm ngành chính
                industries[idx]["la_nganh_chinh"] = True
                self.save_configs()
                self.load_industry_data()
                break

#================================ KHU VỰC QUẢN LÝ XUẤT FILE  ========================================
            
    def preview_word(self):
        self.show_export_popup("Xem Word")

    def export_file(self):
        self.show_export_popup("Xuất file")

    def merge_documents(self, doc_paths, data_lower):
        if not doc_paths or not isinstance(doc_paths, (list, tuple)):
            raise ValueError("doc_paths phải là một danh sách hợp lệ chứa các đường dẫn template")
        
        temp_files = []
        try:
            # Log dữ liệu đầu vào để kiểm tra
            industries = data_lower.get("nganh_nghe", [])
            members = data_lower.get("thanh_vien", [])
            logging.info(f"Dữ liệu đầu vào - Industries: {len(industries)}, Members: {len(members)}")
            logging.info(f"Doc paths: {doc_paths}")

            for i, doc_path in enumerate(doc_paths):
                if not os.path.exists(doc_path):
                    raise FileNotFoundError(f"File không tồn tại: {doc_path}")
                doc = Document(doc_path)

                # Đặt font mặc định cho toàn bộ tài liệu
                style = doc.styles["Normal"]
                font = style.font
                font.name = "Times New Roman"
                font.size = Pt(14)
                font._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            
                # Xử lý tất cả placeholder trong tài liệu
                for paragraph in doc.paragraphs:
                    # Xử lý {{bang_nganh_nghe}}
                    if "{{bang_nganh_nghe}}" in paragraph.text or "{{ bang_nganh_nghe }}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{bang_nganh_nghe}}", "").replace("{{ bang_nganh_nghe }}", "")
                        p = paragraph._element
                        if industries:
                            table = self.create_industry_table(doc, industries)
                            p.getparent().insert(p.getparent().index(p) + 1, table._element)
                        else:
                            p.getparent().insert(p.getparent().index(p) + 1, doc.add_paragraph("Không có thông tin ngành nghề để hiển thị.")._element)

                    # Xử lý {{bang_hop_thanh_vien}}
                    if "{{bang_hop_thanh_vien}}" in paragraph.text or "{{ bang_hop_thanh_vien }}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{bang_hop_thanh_vien}}", "").replace("{{ bang_hop_thanh_vien }}", "")
                        p = paragraph._element
                        if members:
                            table = self.create_member_table(doc, members)
                            if table:
                                p.getparent().insert(p.getparent().index(p) + 1, table._element)
                        else:
                            p.getparent().insert(p.getparent().index(p) + 1, doc.add_paragraph("Không có thông tin thành viên để hiển thị.")._element)

                    # Xử lý {{bang_thanh_vien}}
                    if "{{bang_thanh_vien}}" in paragraph.text or "{{ bang_thanh_vien }}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{bang_thanh_vien}}", "").replace("{{ bang_thanh_vien }}", "")
                        p = paragraph._element
                        if members:
                            table = self.create_member_info_table(doc, members)
                            if table:
                                p.getparent().insert(p.getparent().index(p) + 1, table._element)
                                logging.info("Đã chèn bang_thanh_vien")
                        else:
                            p.getparent().insert(p.getparent().index(p) + 1, doc.add_paragraph("Không có thông tin thành viên để hiển thị.")._element)

                    # Xử lý {{bang_gop_von}}
                    if "{{bang_gop_von}}" in paragraph.text or "{{ bang_gop_von }}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{bang_gop_von}}", "").replace("{{ bang_gop_von }}", "")
                        p = paragraph._element
                        if members:
                            table = self.create_capital_contribution_table(doc, members)
                            if table:
                                p.getparent().insert(p.getparent().index(p) + 1, table._element)
                                logging.info("Đã chèn bang_gop_von")
                        else:
                            p.getparent().insert(p.getparent().index(p) + 1, doc.add_paragraph("Không có thông tin góp vốn để hiển thị.")._element)

                    # Xử lý {{danh_sach_thanh_vien}}
                    if "{{danh_sach_thanh_vien}}" in paragraph.text or "{{ danh_sach_thanh_vien }}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{danh_sach_thanh_vien}}", "").replace("{{ danh_sach_thanh_vien }}", "")
                        p = paragraph._element
                        if members:
                            table = self.create_member_list_table(doc, members)
                            if table:
                                p.getparent().insert(p.getparent().index(p) + 1, table._element)
                                logging.info("Đã chèn danh_sach_thanh_vien")
                        else:
                            p.getparent().insert(p.getparent().index(p) + 1, doc.add_paragraph("Không có danh sách thành viên để hiển thị.")._element)

                # Lưu file tạm thời
                temp_file = f"temp_{i}.docx"
                temp_files.append(temp_file)
                doc.save(temp_file)
                logging.info(f"Đã lưu file tạm thời: {temp_file}")

                # Tạo render_data cho các placeholder còn lại
                render_data = {k: v for k, v in data_lower.items() if k not in ["nganh_nghe", "thanh_vien"]}
                for idx, member in enumerate(members, start=1):
                    for key, value in member.items():
                        render_data[f"{key}_{idx}"] = value
                # Thêm giá trị giả để tránh cảnh báo
                render_data["bang_nganh_nghe"] = ""
                render_data["bang_hop_thanh_vien"] = ""
                render_data["bang_thanh_vien"] = ""
                render_data["bang_gop_von"] = ""
                render_data["danh_sach_thanh_vien"] = ""

                template = DocxTemplate(temp_file)
                template.render(render_data)
                template.save(temp_file)
                logging.info(f"Đã cập nhật file tạm thời với các placeholder còn lại: {temp_file}")

            if not temp_files:
                return None

            if len(temp_files) == 1:  # Chế độ xuất riêng lẻ
                final_doc = Document(temp_files[0])
                logging.info("Hoàn tất chế độ xuất riêng lẻ")
                return final_doc
            else:  # Chế độ xuất gộp
                master = Document(temp_files[0])
                composer = Composer(master)
                for i, temp_file in enumerate(temp_files[1:], 1):
                    doc = Document(temp_file)
                    if i > 0:
                        add_section_break(composer.doc)
                    composer.append(doc)
                logging.info("Hoàn tất chế độ xuất gộp")
                return master

        except Exception as e:
            logging.error(f"Lỗi khi gộp tài liệu: {str(e)}")
            messagebox.showerror("Lỗi Template", f"Lỗi khi xử lý template:\n{str(e)}")
            return None
        finally:
            for temp_file in temp_files:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
        return None
    
    def show_export_popup(self, export_type):
        if not self.current_config_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn cấu hình trước!")
            return
        popup = create_centered_popup(self.root, f"{export_type}", 400, 600)
        popup.bind('<Escape>', lambda e: popup.destroy())
        ttk.Label(popup, text="Chọn template để xuất:").pack(pady=5)
        
        template_frame = ttk.Frame(popup)
        template_frame.pack(fill="both", expand=True, padx=10, pady=5)
        
        canvas = tk.Canvas(template_frame, bg="#ffffff", highlightthickness=0)
        scrollbar = ttk.Scrollbar(template_frame, orient="vertical", command=canvas.yview)
        template_inner_frame = ttk.Frame(canvas)
        template_inner_frame.bind("<Configure>", lambda e, c=canvas: c.configure(scrollregion=c.bbox("all")))
        canvas.create_window((0, 0), window=template_inner_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        templates = self.configs[self.current_config_name].get("templates", {})
        template_vars = {}
        for template in templates.keys():
            var = tk.BooleanVar(value=False)
            template_vars[template] = var
            chk = ttk.Checkbutton(template_inner_frame, text=template, variable=var)
            chk.pack(anchor="w", pady=2)

        export_mode = tk.StringVar(value="merge")
        ttk.Label(popup, text="Chế độ xuất:").pack(pady=5)
        ttk.Radiobutton(popup, text="Gộp thành 1 file", variable=export_mode, value="merge").pack(anchor="w", padx=10)
        ttk.Radiobutton(popup, text="Xuất riêng lẻ", variable=export_mode, value="separate").pack(anchor="w", padx=10)

        file_format = tk.StringVar(value="Word")
        if export_type == "Xuất file":
            ttk.Label(popup, text="Định dạng file:").pack(pady=5)
            ttk.Radiobutton(popup, text="Word (.docx)", variable=file_format, value="Word").pack(anchor="w", padx=10)
            ttk.Radiobutton(popup, text="PDF (.pdf)", variable=file_format, value="PDF").pack(anchor="w", padx=10)

        def confirm_export():
            selected_templates = [t for t, var in template_vars.items() if var.get()]
            if not selected_templates:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn ít nhất một template!")
                return
            popup.destroy()
            data = {field: self.entries[field].get() for field in self.entries}
            selected_name = self.load_data_var.get()
            industries = []
            members = []
            for entry in self.saved_entries:
                if entry["name"] == selected_name:
                    industries = entry["data"].get("nganh_nghe", [])
                    members = entry["data"].get("thanh_vien", [])
                    break
            
            # Chuẩn hóa dữ liệu
            data_lower = {normalize_vietnamese(key): value for key, value in data.items()}
            data_lower["nganh_nghe"] = industries
            data_lower["thanh_vien"] = members
            
            # Tính toán và thêm von_dieu_le_bang_chu nếu von_đieu_le tồn tại
            if "von_dieu_le" in data_lower and data_lower["von_dieu_le"]:
                data_lower["von_dieu_le_bang_chu"] = number_to_words(data_lower["von_dieu_le"])
            
            # Debug: In dữ liệu để kiểm tra
            print("data_lower sau khi thêm von_dieu_le_bang_chu:", data_lower)
            
            doc_paths = [os.path.join(self.templates_dir, t) for t in selected_templates]
            if not all(os.path.exists(dp) for dp in doc_paths):
                missing = [dp for dp in doc_paths if not os.path.exists(dp)]
                messagebox.showerror("Lỗi Template", f"Các file template sau không tồn tại: {', '.join(missing)}")
                return
            
            missing_fields = self.check_template_placeholders(doc_paths, data_lower)
            if missing_fields is None:
                return
            if missing_fields:
                #messagebox.showwarning("Cảnh báo", f"Các trường sau không có trong dữ liệu: {', '.join(missing_fields)}")
                pass
            
            mode = export_mode.get()
            if export_type == "Xem Word":
                self.export_preview(doc_paths, data_lower, mode)
            elif export_type == "Xuất file":
                if file_format.get() == "Word":
                    if mode == "merge":
                        self.export_to_word(doc_paths, data_lower, mode)
                    elif mode == "separate":
                        for doc_path in doc_paths:
                            self.export_to_word([doc_path], data_lower, mode)
                else:
                    self.export_to_pdf(doc_paths, data_lower, mode)

        ttk.Button(popup, text="Xuất File", command=confirm_export, style="primary.TButton").pack(pady=10)

    def export_preview(self, doc_paths, data_lower, mode):
        if mode == "merge":
            # Tính số lượng thành viên
            so_thanh_vien = len(data_lower.get("thanh_vien", []))
            data_lower["so_thanh_vien"] = str(so_thanh_vien)  # Thêm vào data_lower
            
            merged_doc = self.merge_documents(doc_paths, data_lower)
            if merged_doc:
                preview_window = Toplevel(self.root)
                preview_window.title("Xem trước Word")
                preview_window.geometry("800x600")
                text_widget = Text(preview_window, wrap="word")
                text_widget.pack(fill="both", expand=True, padx=10, pady=10)
                for paragraph in merged_doc.paragraphs:
                    text_widget.insert(tk.END, paragraph.text + "\n")
        else:
            messagebox.showinfo("Thông báo", "Chế độ xem trước chỉ hỗ trợ gộp file!")

    def export_to_word(self, doc_paths, data_lower, mode):
        # Tính số lượng thành viên
        so_thanh_vien = len(data_lower.get("thanh_vien", []))
        data_lower["so_thanh_vien"] = str(so_thanh_vien)  # Thêm vào data_lower

        if mode == "merge":
            output_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")],
                initialfile=f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            if output_path:
                merged_doc = self.merge_documents(doc_paths, data_lower)
                if merged_doc:
                    try:
                        merged_doc.save(output_path)
                        os.startfile(output_path)
                        messagebox.showinfo("Thành công", f"Đã xuất file: {output_path}")
                        logging.info(f"Xuất Word gộp: {output_path}")
                    except Exception as e:
                        logging.error(f"Lỗi khi xuất Word: {str(e)}")
                        messagebox.showerror("Lỗi", f"Không thể xuất Word: {str(e)}")
        elif mode == "separate":
            output_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")],
                initialfile=f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                title="Chọn vị trí lưu file đầu tiên"
            )
            if output_path:
                base_path, ext = os.path.splitext(output_path)
                for i, doc_path in enumerate(doc_paths):
                    # Gọi merge_documents cho từng tài liệu riêng lẻ
                    single_doc = self.merge_documents([doc_path], data_lower)
                    if single_doc:
                        file_path = f"{base_path}_{i+1}{ext}"
                        try:
                            single_doc.save(file_path)
                            os.startfile(file_path)
                        except Exception as e:
                            logging.error(f"Lỗi khi lưu hoặc mở file {file_path}: {str(e)}")
                            messagebox.showerror("Lỗi", f"Không thể lưu/mở file {file_path}: {str(e)}")
                messagebox.showinfo("Thành công", f"Đã xuất {len(doc_paths)} file Word riêng lẻ!")
                logging.info(f"Xuất {len(doc_paths)} file Word riêng lẻ từ {base_path}")

    def export_to_pdf(self, doc_paths, data_lower, mode):
        # Tính số lượng thành viên
        so_thanh_vien = len(data_lower.get("thanh_vien", []))
        data_lower["so_thanh_vien"] = str(so_thanh_vien)  # Thêm vào data_lower
    
        output_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")],
                                                  initialfile=f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
        if output_path:
            if mode == "merge":
                merged_doc = self.merge_documents(doc_paths, data_lower)
                if merged_doc:
                    try:
                        temp_docx = "temp_merged.docx"
                        merged_doc.save(temp_docx)
                        convert(temp_docx, output_path)
                        os.remove(temp_docx)
                        os.startfile(output_path)
                        messagebox.showinfo("Thành công", f"Đã xuất file PDF: {output_path}")
                        logging.info(f"Xuất PDF gộp: {output_path}")
                    except Exception as e:
                        logging.error(f"Lỗi khi xuất PDF: {str(e)}")
                        messagebox.showerror("Lỗi", f"Không thể xuất PDF: {str(e)}")
            else:
                base_path, ext = os.path.splitext(output_path)
                for i, doc_path in enumerate(doc_paths):
                    template = DocxTemplate(doc_path)
                    template.render(data_lower)
                    temp_docx = f"temp_{i}.docx"
                    template.save(temp_docx)
                    try:
                        file_path = f"{base_path}_{i+1}.pdf"
                        convert(temp_docx, file_path)
                        os.remove(temp_docx)
                        os.startfile(file_path)
                    except Exception as e:
                        logging.error(f"Lỗi khi chuyển đổi PDF từ {temp_docx}: {str(e)}")
                        messagebox.showerror("Lỗi", f"Không thể chuyển đổi sang PDF: {str(e)}")
                messagebox.showinfo("Thành công", f"Đã xuất {len(doc_paths)} file PDF riêng lẻ!")
                logging.info(f"Xuất {len(doc_paths)} file PDF riêng lẻ từ {base_path}")

    

    


#================================ KHU VỰC QUẢN LÝ CHÈN BẢNG ========================================

    def create_industry_table(self, doc, industries):
        # Tạo bảng với số hàng = 1 (tiêu đề) + số ngành nghề
        table = doc.add_table(rows=1 + len(industries), cols=4)
        table.autofit = False  # Tắt autofit để cố định kích thước

        # Đặt chiều rộng cột
        for cell in table.columns[0].cells:
            cell.width = Cm(1.5)  # Cột 1: STT
        for cell in table.columns[1].cells:
            cell.width = Cm(9.5)  # Cột 2: Tên ngành
        for cell in table.columns[2].cells:
            cell.width = Cm(2.75)  # Cột 3: Mã ngành
        for cell in table.columns[3].cells:
            cell.width = Cm(3.0)  # Cột 4: Ngành nghề kinh doanh chính

        # Thêm viền cho bảng
        table_style = table._tblPr.xpath('./w:tblBorders')[0] if table._tblPr.xpath('./w:tblBorders') else OxmlElement('w:tblBorders')
        for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            table_style.append(border)
        if not table._tblPr.xpath('./w:tblBorders'):
            table._tblPr.append(table_style)

        # Đặt tiêu đề cho các cột
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "STT"
        hdr_cells[1].text = "Tên ngành"
        hdr_cells[2].text = "Mã ngành"
        hdr_cells[3].text = "Ngành nghề kinh doanh chính"

        # Định dạng chữ cho tiêu đề
        for i, cell in enumerate(hdr_cells):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.text = cell.text
            font = run.font
            font.size = Pt(14)
            font.bold = True

        # Thêm dữ liệu ngành nghề
        for idx, industry in enumerate(industries, start=1):
            row_cells = table.rows[idx].cells  # Lấy hàng đã tạo sẵn
            row_cells[0].text = str(idx)
            row_cells[1].text = industry.get("ten_nganh", "")
            row_cells[2].text = industry.get("ma_nganh", "")
            row_cells[3].text = "X" if industry.get("la_nganh_chinh", False) else ""

            # Định dạng chữ cho dữ liệu
            for i, cell in enumerate(row_cells):
                paragraph = cell.paragraphs[0]
                if i in [0, 2, 3]:
                    paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = cell.text
                font = run.font
                font.size = Pt(14)

        # Căn giữa bảng
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        return table

    def create_member_table(self, doc, members):
        """Tạo bảng họp thành viên (bang_hop_thanh_vien)."""
        if not members:
            p = doc.add_paragraph()
            p.text = "Không có thông tin thành viên để hiển thị."
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return None

        # Tạo bảng với 2 cột
        table = doc.add_table(rows=len(members), cols=2)
        table.autofit = True  # bật chế độ tự động điều chỉnh độ rộng

        # Đặt chiều rộng cột
        widths = [5.5, 11.5] 
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width) 

        # Điền dữ liệu vào bảng
        for i, member in enumerate(members, start=1):  
            so_giay_chung_nhan = f"{i:02d}/GCN"  # Tạo số giấy chứng nhận

            # Xác định vai trò dựa trên la_chu_tich
            if member.get("la_chu_tich", False):
                role = "Chủ tịch hội đồng thành viên – Chủ toạ cuộc họp sở hữu"
            else:
                role = "Thành viên góp vốn sở hữu"

            # Cột 1: Tên thành viên
            cell_1 = table.cell(i-1, 0)
            p1 = cell_1.paragraphs[0]
            p1.text = member.get("ho_ten", "")
            p1.style = doc.styles["Normal"]
            p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run1 = p1.runs[0]
            run1.font.name = "Times New Roman"
            run1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run1.font.size = Pt(14)

            # Cột 2: Thông tin
            cell_2 = table.cell(i-1, 1)
            p2 = cell_2.paragraphs[0]
            p2.text = (
                f"– {role} {member.get('von_gop', '')} "
                f"chiếm tỷ lệ {member.get('ty_le_gop', '')}% vốn điều lệ.\n"
                f"Giấy chứng nhận góp vốn số {so_giay_chung_nhan}, cấp ngày {member.get('ngay_gop_von', '')}.\n"
            )
            p2.style = doc.styles["Normal"]
            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run2 = p2.runs[0]
            run2.font.name = "Times New Roman"
            run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run2.font.size = Pt(14)

        return table
    
    def create_member_info_table(self, doc, members):
        """Tạo bảng thông tin thành viên (bang_thanh_vien) với 7 cột."""
        if not members:
            p = doc.add_paragraph()
            p.text = "Không có thông tin thành viên để hiển thị."
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return None

        # Tạo bảng với tiêu đề và dữ liệu
        table = doc.add_table(rows=1 + len(members), cols=7)
        table.autofit = False

        # Đặt chiều rộng cột
        widths = [1.5, 3.5, 2.75, 1.5, 1.5, 4.25, 3.75]  # 7 cột
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

        # Thêm viền cho bảng
        table_style = table._tblPr.xpath('./w:tblBorders')[0] if table._tblPr.xpath('./w:tblBorders') else OxmlElement('w:tblBorders')
        for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            table_style.append(border)
        if not table._tblPr.xpath('./w:tblBorders'):
            table._tblPr.append(table_style)

        # Thêm tiêu đề
        hdr_cells = table.rows[0].cells
        headers = [
            "STT", 
            "Tên thành viên", 
            "Ngày, tháng, năm sinh đối với thành viên là cá nhân", 
            "Giới tính", 
            "Quốc tịch", 
            "Địa chỉ liên lạc đối với cá nhân, hoặc địa chỉ trụ sở chính đối với tổ chức", 
            "Loại giấy tờ, số, ngày cấp, cơ quan cấp Giấy tờ pháp lý của cá nhân/tổ chức"
        ]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            paragraph = hdr_cells[i].paragraphs[0]
            paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.text = header
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(13)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        # Điền dữ liệu
        for idx, member in enumerate(members, start=1):
            row_cells = table.rows[idx].cells
            row_cells[0].text = str(idx)
            row_cells[1].text = member.get("ho_ten", "")
            row_cells[2].text = member.get("ngay_sinh", "")
            row_cells[3].text = member.get("gioi_tinh", "")
            row_cells[4].text = member.get("quoc_tich", "")
            row_cells[5].text = member.get("dia_chi_lien_lac", "")
            row_cells[6].text = (
                f"{member.get('loai_giay_to', '')}\n"
                f"{member.get('so_cccd', '')}\n"
                f"{member.get('ngay_cap', '')}\n"
                f"{member.get('noi_cap', '')}"
            )

            # Định dạng các ô
            for i, cell in enumerate(row_cells):
                paragraph = cell.paragraphs[0]
                paragraph.alignment = (
                    WD_TABLE_ALIGNMENT.CENTER if i in [0, 3, 4] else WD_ALIGN_PARAGRAPH.CENTER
                )
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = cell.text
                font = run.font
                font.size = Pt(13)
                font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")  # Đảm bảo font cho tiếng Việt

        # Căn giữa bảng
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        return table

    def create_capital_contribution_table(self, doc, members):
        """Tạo bảng góp vốn (bang_gop_von) với 7 cột, gộp cột Vốn góp và thêm cột Ghi chú."""
        logging.info(f"Số lượng thành viên trong create_capital_contribution_table: {len(members)}")
        logging.info(f"Dữ liệu thành viên: {members}")
        
        if not members:
            p = doc.add_paragraph()
            p.text = "Không có danh sách thành viên để hiển thị."
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return None

        # Tạo bảng với 2 hàng tiêu đề (tiêu đề chính gộp, tiêu đề phụ) + số hàng dữ liệu
        table = doc.add_table(rows=2 + len(members), cols=7)
        table.autofit = False

        # Đặt chiều rộng cột
        widths = [1.27, 3.5, 4.41, 1.75, 2.25, 3.5, 1.25]  # 7 cột, bao gồm cột Ghi chú
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

        # Thêm viền cho bảng
        table_style = table._tblPr.xpath('./w:tblBorders')[0] if table._tblPr.xpath('./w:tblBorders') else OxmlElement('w:tblBorders')
        for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            table_style.append(border)
        if not table._tblPr.xpath('./w:tblBorders'):
            table._tblPr.append(table_style)
        
        # Hàng 1: Tiêu đề chính
        main_hdr_cells = table.rows[0].cells
        main_headers = [
            "STT", "Tên thành viên", "Vốn góp", "", "", "Thời hạn góp vốn", "Ghi chú"
        ]
        if len(main_headers) != 7:
            logging.error(f"Số lượng tiêu đề chính không khớp: {len(main_headers)} (yêu cầu 7)")
            raise ValueError("Số lượng tiêu đề chính không khớp với số cột")

        # Gộp hàng 1 và hàng 2 cho cột 1-2 và 6, đồng thời điền tiêu đề chính
        for i in range(len(main_headers)):
            if i in [0, 1, 5, 6]:  # Cột 1-2 (chỉ số 0-1), 6 (chỉ số 5), và 7 (chỉ số 6 - Ghi chú)
                # Gộp ô từ hàng 1 và hàng 2
                cell = main_hdr_cells[i]
                cell.merge(table.rows[1].cells[i])  # Gộp ô hàng 1 và hàng 2
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.add_run(main_headers[i])
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            elif i == 2:  # Cột 3: "Vốn góp"
                cell = main_hdr_cells[2]
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.add_run("Vốn góp")
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                # Gộp cột 3, 4, 5 (chỉ số 2, 3, 4) ở hàng 1
                cell.merge(main_hdr_cells[3])
                cell.merge(main_hdr_cells[4])
            elif i in [3, 4]:
                continue  # Bỏ qua vì đã gộp vào cột 3

        # Hàng 2: Tiêu đề phụ (chỉ hiển thị cho các cột con của "Vốn góp", các cột khác đã gộp)
        sub_hdr_cells = table.rows[1].cells
        for i in range(7):
            if i in [2, 3, 4]:  # Chỉ hiển thị tiêu đề phụ cho cột 3, 4, 5
                sub_headers = ["Phần vốn góp (VNĐ)", "Tỷ lệ (%)", "Loại tài sản, số lượng, giá trị tài sản góp vốn"]
                sub_hdr_cells[i].text = sub_headers[i - 2]
                paragraph = sub_hdr_cells[i].paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = sub_headers[i - 2]
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        # Điền dữ liệu
        for idx, member in enumerate(members, start=1):
            row_cells = table.rows[idx + 1].cells  # +1 vì có 2 hàng tiêu đề
            row_cells[0].text = str(idx)
            row_cells[1].text = member.get("ho_ten", "")
            row_cells[2].text = member.get("von_gop", "")
            row_cells[3].text = member.get("ty_le_gop", "")
            row_cells[4].text = "Đồng Việt Nam"  # Khớp với tài liệu
            row_cells[5].text = member.get("ngay_gop_von", "")
            row_cells[6].text = ""  # Ghi chú để trống

            for i, cell in enumerate(row_cells):
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = cell.text
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        logging.info("Đã chèn bang_gop_von")
        return table

    def create_member_list_table(self, doc, members):
        """Tạo bảng danh sách thành viên (danh_sach_thanh_vien) với 14 cột và gộp cột Vốn góp."""
        logging.info(f"Số lượng thành viên trong create_member_list_table: {len(members)}")
        logging.info(f"Dữ liệu thành viên: {members}")
        
        if not members:
            p = doc.add_paragraph()
            p.text = "Không có danh sách thành viên để hiển thị."
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return None

        # Tạo bảng với 3 hàng tiêu đề (tiêu đề chính gộp, tiêu đề phụ, số thứ tự cột) + số hàng dữ liệu
        table = doc.add_table(rows=3 + len(members), cols=14)
        table.autofit = False

        # Đặt chiều rộng cột
        widths = [1.2, 3.05, 1.62, 1.38, 1.5, 1.5, 3.4, 3.35, 2.75, 1.25, 1.5, 2, 2.5, 1.09]  # Thêm cột Ghi chú
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

        # Thêm viền cho bảng
        table_style = table._tblPr.xpath('./w:tblBorders')[0] if table._tblPr.xpath('./w:tblBorders') else OxmlElement('w:tblBorders')
        for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            table_style.append(border)
        if not table._tblPr.xpath('./w:tblBorders'):
            table._tblPr.append(table_style)
        
        # Hàng 1: Tiêu đề chính
        main_hdr_cells = table.rows[0].cells
        main_headers = [
            "STT", "Tên thành viên", "Ngày, tháng, năm sinh đối với thành viên là cá nhân", "Giới tính", "Quốc tịch", "Dân tộc",
            "Địa chỉ liên lạc đối với thành viên là cá nhân; địa chỉ trụ sở chính đối với thành viên là tổ chức",
            "Loại giấy tờ, số, ngày cấp, cơ quan cấp Giấy tờ pháp lý của cá nhân/tổ chức", "Vốn góp", "", "", "Thời hạn góp vốn",
            "Chữ ký của thành viên", "Ghi chú"
        ]
        if len(main_headers) != 14:
            logging.error(f"Số lượng tiêu đề chính không khớp: {len(main_headers)} (yêu cầu 14)")
            raise ValueError("Số lượng tiêu đề chính không khớp với số cột")

        # Gộp hàng 1 và hàng 2 cho cột 1-8 và 12-14, đồng thời điền tiêu đề chính
        for i in range(len(main_headers)):
            if i in [0, 1, 2, 3, 4, 5, 6, 7, 11, 12, 13]:  # Cột 1-8 (chỉ số 0-7) và 12-14 (chỉ số 11-13)
                # Gộp ô từ hàng 1 và hàng 2
                cell = main_hdr_cells[i]
                cell.merge(table.rows[1].cells[i])  # Gộp ô hàng 1 và hàng 2
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.add_run(main_headers[i])
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            elif i == 8:  # Cột 9: "Vốn góp"
                cell = main_hdr_cells[8]
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.add_run("Vốn góp")
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                # Gộp cột 9, 10, 11 (chỉ số 8, 9, 10) ở hàng 1
                cell.merge(main_hdr_cells[9])
                cell.merge(main_hdr_cells[10])
            elif i in [9, 10]:
                continue  # Bỏ qua vì đã gộp vào cột 9

        # Hàng 2: Tiêu đề phụ (chỉ hiển thị cho các cột con của "Vốn góp", các cột khác đã gộp)
        sub_hdr_cells = table.rows[1].cells
        for i in range(14):
            if i in [8, 9, 10]:  # Chỉ hiển thị tiêu đề phụ cho cột 9, 10, 11
                sub_headers = ["Phần vốn góp (VNĐ)", "Tỷ lệ (%)", "Loại tài sản, số lượng, giá trị tài sản góp vốn"]
                sub_hdr_cells[i].text = sub_headers[i - 8]
                paragraph = sub_hdr_cells[i].paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = sub_headers[i - 8]
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        # Hàng 3: Số thứ tự cột (1, 2, 3, ..., 14)
        num_hdr_cells = table.rows[2].cells
        for i in range(14):
            num_hdr_cells[i].text = str(i + 1)
            paragraph = num_hdr_cells[i].paragraphs[0]
            paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.text = str(i + 1)
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(12)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        # Điền dữ liệu
        for idx, member in enumerate(members, start=1):
            row_cells = table.rows[idx + 2].cells  # +2 vì có 3 hàng tiêu đề
            row_cells[0].text = str(idx)
            row_cells[1].text = member.get("ho_ten", "")
            row_cells[2].text = member.get("ngay_sinh", "")
            row_cells[3].text = member.get("gioi_tinh", "")
            row_cells[4].text = member.get("quoc_tich", "")
            row_cells[5].text = member.get("dan_toc", "")
            row_cells[6].text = member.get("dia_chi_lien_lac", "")
            row_cells[7].text = f"{member.get('loai_giay_to', '')}\n{member.get('so_cccd', '')}\n{member.get('ngay_cap', '')}\n{member.get('noi_cap', '')}"
            row_cells[8].text = member.get("von_gop", "")
            row_cells[9].text = member.get("ty_le_gop", "")
            row_cells[10].text = "Đồng Việt Nam"  # Khớp với tài liệu
            row_cells[11].text = member.get("ngay_gop_von", "")
            row_cells[12].text = ""  # Chữ ký để trống
            row_cells[13].text = ""  # Ghi chú để trống

            for i, cell in enumerate(row_cells):
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = cell.text
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        logging.info("Đã chèn danh_sach_thanh_vien")
        return table    

#================================ KHU VỰC QUẢN LÝ BACKUP ========================================

    def auto_backup(self):
        """Sao lưu tự động dữ liệu mỗi 5 phút vào thư mục backup."""
        try:
            backup_file = os.path.join(self.backup_dir, f"backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
            with open(backup_file, 'w', encoding='utf-8') as f:
                json.dump(self.configs, f, ensure_ascii=False, indent=4)
            logging.info(f"Sao lưu tự động: {backup_file}")
            self.root.after(600000, self.auto_backup)
        except Exception as e:
            logging.error(f"Lỗi khi sao lưu tự động: {str(e)}")

    def restore_from_backup(self):
        """Hiển thị popup để chọn và khôi phục dữ liệu từ file sao lưu, với tùy chọn xóa file cũ."""
        if not os.path.exists(self.backup_dir) or not os.listdir(self.backup_dir):
            messagebox.showinfo("Thông báo", "Chưa có file sao lưu nào!")
            return

        popup = create_centered_popup(self.root, "Khôi phục từ sao lưu", 500, 500)
        popup.bind('<Escape>', lambda e: popup.destroy())
        ttk.Label(popup, text="Chọn file sao lưu để khôi phục:").pack(pady=5)

        backup_frame = ttk.Frame(popup)
        backup_frame.pack(fill="both", expand=True, padx=10, pady=5)
        
        backup_tree = ttk.Treeview(backup_frame, columns=("lastmod"), show="tree headings", height=10, selectmode="browse")
        backup_tree.heading("#0", text="Tên file")
        backup_tree.heading("lastmod", text="Ngày tạo")
        backup_tree.column("#0", width=250)
        backup_tree.column("lastmod", width=150)
        backup_tree.pack(side="left", fill="both", expand=True)
        scrollbar = ttk.Scrollbar(backup_frame, orient="vertical", command=backup_tree.yview)
        scrollbar.pack(side="right", fill="y")
        backup_tree.configure(yscrollcommand=scrollbar.set)

        backup_files = sorted([f for f in os.listdir(self.backup_dir) if f.endswith(".json")], reverse=True)
        for backup_file in backup_files:
            file_path = os.path.join(self.backup_dir, backup_file)
            timestamp = datetime.fromtimestamp(os.path.getctime(file_path)).strftime("%Y-%m-%d %H:%M:%S")
            backup_tree.insert("", "end", text=backup_file, values=(timestamp,))

        def confirm_restore():
            selected = backup_tree.selection()
            if not selected:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn một file sao lưu!")
                return
            backup_file = backup_tree.item(selected[0])["text"]
            file_path = os.path.join(self.backup_dir, backup_file)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    self.configs = json.load(f)
                self.save_configs()
                self.load_selected_config(None)
                popup.destroy()
                messagebox.showinfo("Thành công", f"Đã khôi phục từ file: {backup_file}")
                logging.info(f"Khôi phục từ sao lưu: {backup_file}")
            except Exception as e:
                logging.error(f"Lỗi khi khôi phục: {str(e)}")
                messagebox.showerror("Lỗi", f"Không thể khôi phục: {str(e)}")

        def delete_old_backups():
            if messagebox.askyesno("Xác nhận", "Bạn có muốn xóa các file sao lưu cũ (giữ lại 10 file mới nhất)?"):
                backup_files_sorted = sorted(
                    [f for f in os.listdir(self.backup_dir) if f.endswith(".json")],
                    key=lambda x: os.path.getctime(os.path.join(self.backup_dir, x)),
                    reverse=True
                )
                files_to_delete = backup_files_sorted[10:]
                for file in files_to_delete:
                    os.remove(os.path.join(self.backup_dir, file))
                messagebox.showinfo("Thành công", f"Đã xóa {len(files_to_delete)} file sao lưu cũ!")
                logging.info(f"Xóa {len(files_to_delete)} file sao lưu cũ")
                for item in backup_tree.get_children():
                    backup_tree.delete(item)
                remaining_files = sorted([f for f in os.listdir(self.backup_dir) if f.endswith(".json")], reverse=True)
                for backup_file in remaining_files:
                    file_path = os.path.join(self.backup_dir, backup_file)
                    timestamp = datetime.fromtimestamp(os.path.getctime(file_path)).strftime("%Y-%m-%d %H:%M:%S")
                    backup_tree.insert("", "end", text=backup_file, values=(timestamp,))

        button_frame = ttk.Frame(popup)
        button_frame.pack(pady=10)
        ttk.Button(button_frame, text="Khôi phục", command=confirm_restore, style="primary.TButton").pack(side="left", padx=5)
        ttk.Button(button_frame, text="Xóa file cũ", command=delete_old_backups, style="danger.TButton").pack(side="left", padx=5)
  
if __name__ == "__main__":
    root = TkinterDnD.Tk()
    app = DataEntryApp(root)
    root.mainloop()
