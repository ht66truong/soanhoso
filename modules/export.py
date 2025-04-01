# Thư viện chuẩn
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, Text
import os
import json
import logging
from datetime import datetime
from docx.oxml.ns import qn
from docx2pdf import convert 
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxcompose.composer import Composer
from docx.oxml import OxmlElement

# Thư viện bên thứ ba
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
import ttkbootstrap as ttk
from ttkbootstrap.constants import *

# Module nội bộ
from modules.utils import (create_popup, number_to_words, normalize_vietnamese,
                            add_section_break)


class ExportManager:
    def __init__(self, app):
        self.app = app

    def export_data(self):
        """Xuất dữ liệu ra file Excel hoặc JSON."""
        # Tạo popup
        popup = create_popup (self.app.root, "Chọn định dạng xuất", 300, 150)
        ttk.Label(popup, text="Chọn định dạng:").pack(pady=10)
        file_format = tk.StringVar(value="Excel")
        ttk.Radiobutton(popup, text="Excel (.xlsx)", variable=file_format, value="Excel").pack(anchor="w", padx=20)
        ttk.Radiobutton(popup, text="JSON (.json)", variable=file_format, value="JSON").pack(anchor="w", padx=20)

        def confirm_export():
            popup.destroy()
            if file_format.get() == "Excel":
                output_path = filedialog.asksaveasfilename(
                    defaultextension=".xlsx",
                    filetypes=[("Excel files", "*.xlsx")],
                    initialfile=f"data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                )
                if not output_path:
                    return
                data = {field: self.app.entries[field].get() for field in self.app.entries}
                selected_name = self.app.load_data_var.get()
                for entry in self.app.saved_entries:
                    if entry["name"] == selected_name:
                        industries = entry["data"].get("nganh_nghe", [])
                        members = entry["data"].get("thanh_vien", [])
                        df_main = pd.DataFrame([data])
                        df_industries = pd.DataFrame(industries)
                        df_members = pd.DataFrame(members)
                        break
                else:
                    df_main = pd.DataFrame([data])
                    df_industries = pd.DataFrame()
                    df_members = pd.DataFrame()

                try:
                    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
                        df_main.to_excel(writer, sheet_name="Main", index=False)
                        df_industries.to_excel(writer, sheet_name="Industries", index=False)
                        df_members.to_excel(writer, sheet_name="Members", index=False)
                    os.startfile(output_path)
                    messagebox.showinfo("Thành công", f"Đã xuất file Excel: {output_path}")
                    logging.info(f"Xuất Excel: {output_path}")
                except Exception as e:
                    logging.error(f"Lỗi khi xuất Excel: {str(e)}")
                    messagebox.showerror("Lỗi", f"Không thể xuất file Excel: {str(e)}")

            elif file_format.get() == "JSON":
                output_path = filedialog.asksaveasfilename(
                    defaultextension=".json",
                    filetypes=[("JSON files", "*.json")],
                    initialfile=f"data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                )
                if not output_path:
                    return
                data = {field: self.app.entries[field].get() for field in self.app.entries}
                selected_name = self.app.load_data_var.get()
                export_data = None
                for entry in self.app.saved_entries:
                    if entry["name"] == selected_name:
                        export_data = {
                            "name": selected_name,
                            "data": {
                                **data,
                                "nganh_nghe": entry["data"].get("nganh_nghe", []),
                                "thanh_vien": entry["data"].get("thanh_vien", [])
                            }
                        }
                        break
                if export_data is None:
                    export_data = {
                        "name": selected_name,
                        "data": {
                            **data,
                            "nganh_nghe": [],
                            "thanh_vien": []
                        }
                    }

                try:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        json.dump(export_data, f, ensure_ascii=False, indent=4)
                    os.startfile(output_path)
                    messagebox.showinfo("Thành công", f"Đã xuất file JSON: {output_path}")
                    logging.info(f"Xuất JSON: {output_path}")
                except Exception as e:
                    logging.error(f"Lỗi khi xuất JSON: {str(e)}")
                    messagebox.showerror("Lỗi", f"Không thể xuất file JSON: {str(e)}")

        ttk.Button(popup, text="Xuất", command=confirm_export, style="primary.TButton").pack(pady=10)

    def import_from_file(self):
        """Nhập dữ liệu từ file Excel hoặc JSON."""
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx"), ("JSON files", "*.json")])
        if not file_path:
            return
        try:
            if file_path.endswith(".xlsx"):
                # Đọc sheet "Main" cho dữ liệu chính
                df_main = pd.read_excel(file_path, sheet_name="Main", dtype=str)  # Ép tất cả dữ liệu thành chuỗi
                # Thay thế NaN bằng chuỗi rỗng trong dữ liệu chính
                data = df_main.iloc[0].fillna("").to_dict()
                for field in self.app.entries:
                    if field in data:
                        self.app.entries[field].delete(0, tk.END)
                        self.app.entries[field].insert(0, data[field])  # Giữ nguyên chuỗi, bao gồm số 0 ở đầu

                # Đọc sheet "Industries" cho dữ liệu ngành nghề
                try:
                    df_industries = pd.read_excel(file_path, sheet_name="Industries", dtype=str)  # Ép dữ liệu thành chuỗi
                    # Thay thế NaN bằng chuỗi rỗng trong dữ liệu ngành nghề
                    industries = df_industries.fillna("").to_dict(orient="records")
                except ValueError:
                    industries = []

                # Đọc sheet "Members" cho dữ liệu thành viên
                try:
                    df_members = pd.read_excel(file_path, sheet_name="Members", dtype=str)  # Ép dữ liệu thành chuỗi
                    # Thay thế NaN bằng chuỗi rỗng trong dữ liệu thành viên
                    members = df_members.fillna("").to_dict(orient="records")
                except ValueError:
                    members = []

                selected_name = self.app.load_data_var.get()
                for entry in self.app.saved_entries:
                    if entry["name"] == selected_name:
                        entry["data"] = {field: self.app.entries[field].get() for field in self.app.entries}
                        entry["data"]["nganh_nghe"] = industries
                        entry["data"]["thanh_vien"] = members
                        break
                else:
                    new_entry = {
                        "name": selected_name,
                        "data": {field: self.app.entries[field].get() for field in self.app.entries}
                    }
                    new_entry["data"]["nganh_nghe"] = industries
                    new_entry["data"]["thanh_vien"] = members
                    self.app.saved_entries.append(new_entry)
                    self.app.load_data_dropdown["values"] = [entry["name"] for entry in self.app.saved_entries]

                self.app.config_manager.save_configs()
                self.app.member_manager.load_member_data()
                self.app.industry_manager.load_industry_data()
                messagebox.showinfo("Thành công", "Đã nhập dữ liệu từ file Excel!")
                logging.info(f"Nhập dữ liệu từ Excel: {file_path}")
            elif file_path.endswith(".json"):
                with open(file_path, 'r', encoding='utf-8') as f:
                    imported_data = json.load(f)
                if isinstance(imported_data, dict) and "name" in imported_data and "data" in imported_data:
                    self.app.saved_entries.append(imported_data)
                    self.app.load_data_dropdown["values"] = [entry["name"] for entry in self.app.saved_entries]
                    self.app.load_data_var.set(imported_data["name"])
                    self.app.load_selected_entry(None)
                    self.app.config_manager.save_configs()
                    messagebox.showinfo("Thành công", "Đã nhập dữ liệu từ file JSON!")
                    logging.info(f"Nhập dữ liệu từ JSON: {file_path}")
        except Exception as e:
            logging.error(f"Lỗi khi nhập file: {str(e)}")
            messagebox.showerror("Lỗi", f"Không thể nhập file: {str(e)}")
            
    def show_placeholder_popup(self):
        """Hiển thị popup chứa danh sách placeholder và nút xuất."""
        if not self.app.config_manager.current_config_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn cấu hình trước!")
            return

        popup = create_popup(self.app.root, "Danh sách Placeholder", 600, 400)
        placeholder_frame = ttk.Frame(popup)
        placeholder_frame.pack(fill="both", expand=True, padx=10, pady=10)
        placeholder_list = ttk.Treeview(placeholder_frame, columns=("original"), show="tree headings", height=10)
        placeholder_list.heading("#0", text="Placeholder")
        placeholder_list.heading("original", text="Tên trường gốc")
        placeholder_list.column("#0", width=250)
        placeholder_list.column("original", width=250)

        scrollbar = ttk.Scrollbar(placeholder_frame, orient="vertical", command=placeholder_list.yview)
        scrollbar.pack(side="right", fill="y")
        placeholder_list.configure(yscrollcommand=scrollbar.set)
        placeholder_list.pack(side="left", fill="both", expand=True)

        placeholders = {normalize_vietnamese(field): field for field in self.app.fields}
        for placeholder, original in sorted(placeholders.items()):
            placeholder_list.insert("", "end", text=f"{{{{ {placeholder} }}}}", values=(original,))

        context_menu = tk.Menu(popup, tearoff=0)
        context_menu.add_command(label="Sao chép Placeholder", command=lambda: self.copy_placeholder_from_popup(placeholder_list))

        def show_context_menu(event):
            item = placeholder_list.identify_row(event.y)
            if item:
                placeholder_list.selection_set(item)
                context_menu.post(event.x_root, event.y_root)

        placeholder_list.bind("<Button-3>", show_context_menu)
        placeholder_list.bind("<Double-1>", lambda event: self.copy_placeholder_from_popup(placeholder_list))
        placeholder_list.bind("<Control-c>", lambda event: self.copy_placeholder_from_popup(placeholder_list))

        button_frame = ttk.Frame(popup)
        button_frame.pack(fill="x", pady=5)
        ttk.Button(button_frame, text="Xuất Placeholder", command=self.export_placeholders, style="primary.TButton").pack(side="left", padx=5)
        ttk.Button(button_frame, text="Đóng", command=popup.destroy, style="danger.TButton").pack(side="right", padx=5)

    def copy_placeholder_from_popup(self, placeholder_list):
        selected = placeholder_list.selection()
        if selected:
            placeholder = placeholder_list.item(selected[0])["text"]
            self.app.root.clipboard_clear()
            self.app.root.clipboard_append(placeholder)
            
    def export_placeholders(self):
        """Xuất danh sách placeholder ra file Word hoặc text."""
        if not self.app.config_manager.current_config_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn cấu hình trước!")
            return

        placeholders = {}
        for field in self.app.fields:
            normalized = normalize_vietnamese(field)
            placeholders[normalized] = field

        popup = create_popup(self.app.root, "Xuất Placeholder", 300, 150) 
        ttk.Label(popup, text="Chọn định dạng xuất:").pack(pady=10)
        file_format = tk.StringVar(value="Word")
        ttk.Radiobutton(popup, text="Word (.docx)", variable=file_format, value="Word").pack(anchor="w", padx=20)
        ttk.Radiobutton(popup, text="Text (.txt)", variable=file_format, value="Text").pack(anchor="w", padx=20)

        def confirm_export():
            output_path = filedialog.asksaveasfilename(
                defaultextension=f".{file_format.get().lower()}",
                filetypes=[("Word files", "*.docx") if file_format.get() == "Word" else ("Text files", "*.txt")],
                initialfile=f"placeholders_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{file_format.get().lower()}"
            )
            if not output_path:
                popup.destroy()
                return

            try:
                if file_format.get() == "Word":
                    doc = Document()
                    doc.add_heading("Danh sách Placeholder", level=1)
                    doc.add_paragraph("Sao chép các placeholder dưới đây và dán vào template Word theo cú pháp {{ placeholder }}:")
                    table = doc.add_table(rows=1, cols=2)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "Placeholder"
                    hdr_cells[1].text = "Tên trường gốc"
                    for placeholder, original in sorted(placeholders.items()):
                        row_cells = table.add_row().cells
                        row_cells[0].text = f"{{{{ {placeholder} }}}}"
                        row_cells[1].text = original
                    doc.save(output_path)
                else:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write("Danh sách Placeholder\n")
                        f.write("Sao chép các placeholder dưới đây và dán vào template Word theo cú pháp {{ placeholder }}:\n\n")
                        f.write("Placeholder\tTên trường gốc\n")
                        f.write("-" * 50 + "\n")
                        for placeholder, original in sorted(placeholders.items()):
                            f.write(f"{{{{ {placeholder} }}}}\t{original}\n")
                os.startfile(output_path)
                messagebox.showinfo("Thành công", f"Đã xuất danh sách placeholder ra: {output_path}")
                logging.info(f"Xuất danh sách placeholder: {output_path}")
                popup.destroy()
            except Exception as e:
                logging.error(f"Lỗi khi xuất placeholder: {str(e)}")
                messagebox.showerror("Lỗi", f"Không thể xuất file: {str(e)}")
                popup.destroy()

    def check_template_placeholders(self, doc_paths, data_lower):
        try:
            placeholders = set()
            for doc_path in doc_paths:
                doc = DocxTemplate(doc_path)
                for item in doc.get_undeclared_template_variables():
                    # Bỏ qua các placeholder liên quan đến thành viên (dạng ho_ten_1, von_gop_1, v.v.)
                    if any(field in item for field in self.app.member_columns) and any(item.endswith(f"_{i}") for i in range(1, 100)):
                        continue
                    placeholders.add(item)
            missing_fields = [p for p in placeholders if p not in data_lower]
            return missing_fields
        except Exception as e:
            logging.error(f"Lỗi khi kiểm tra placeholder: {str(e)}")
            messagebox.showerror("Lỗi Template", f"Lỗi khi kiểm tra placeholder:\n{str(e)}")
            return None

    def preview_word(self):
        self.show_export_popup("Xem Word")

    def export_file(self):
        self.show_export_popup("Xuất file")

    def show_export_popup(self, export_type):
        if not self.app.config_manager.current_config_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn cấu hình trước!")
            return
        popup = create_popup(self.app.root, f"{export_type}", 400, 600)
        ttk.Label(popup, text="Chọn template để xuất:").pack(pady=5)
        
        template_frame = ttk.Frame(popup)
        template_frame.pack(fill="both", expand=True, padx=10, pady=5)
        
        canvas = tk.Canvas(template_frame, bg="#ffffff", highlightthickness=0)
        scrollbar = ttk.Scrollbar(template_frame, orient="vertical", command=canvas.yview)
        template_inner_frame = ttk.Frame(canvas)
        template_inner_frame.bind("<Configure>", lambda e, c=canvas: c.configure(scrollregion=c.bbox("all")))
        canvas.create_window((0, 0), window=template_inner_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        templates = self.app.config_manager.configs[self.app.config_manager.current_config_name].get("templates", {})
        template_vars = {}
        for template in templates.keys():
            # Loại bỏ đuôi .docx khỏi tên hiển thị
            display_name = os.path.splitext(template)[0]
            var = tk.BooleanVar(value=False)
            template_vars[template] = var
            chk = ttk.Checkbutton(template_inner_frame, text=display_name, variable=var)
            chk.pack(anchor="w", pady=2)

        export_mode = tk.StringVar(value="separate")
        ttk.Label(popup, text="Chế độ xuất:").pack(pady=5)
        ttk.Radiobutton(popup, text="Xuất riêng lẻ", variable=export_mode, value="separate").pack(anchor="w", padx=10)
        ttk.Radiobutton(popup, text="Xuất gộp thành 1 file", variable=export_mode, value="merge").pack(anchor="w", padx=10)

        file_format = tk.StringVar(value="Word")
        if export_type == "Xuất file":
            ttk.Label(popup, text="Định dạng file:").pack(pady=5)
            ttk.Radiobutton(popup, text="Word (.docx)", variable=file_format, value="Word").pack(anchor="w", padx=10)
            ttk.Radiobutton(popup, text="PDF (.pdf)", variable=file_format, value="PDF").pack(anchor="w", padx=10)

        # Thêm trường chọn ngày tháng năm
        date_frame = ttk.Frame(popup)  # Tạo một frame để chứa nhãn và trường nhập liệu
        date_frame.pack(pady=10)  # Căn giữa frame trong popup

        ttk.Label(date_frame, text="Chọn ngày tháng năm:\n     (dd/mm/yyyy)", anchor="center").pack(pady=5)  # Nhãn căn giữa
        date_var = tk.StringVar(value=datetime.now().strftime("%d/%m/%Y"))  # Mặc định là ngày hiện tại
        date_entry = ttk.Entry(date_frame, textvariable=date_var, width=20, justify="center")  # Trường nhập liệu căn giữa
        date_entry.pack(pady=5)

        def confirm_export():
            selected_templates = [t for t, var in template_vars.items() if var.get()]
            if not selected_templates:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn ít nhất một template!")
                return
            popup.destroy()
            data = {field: self.app.entries[field].get() for field in self.app.entries}
            selected_name = self.app.load_data_var.get()
            industries = []
            members = []
            for entry in self.app.saved_entries:
                if entry["name"] == selected_name:
                    industries = entry["data"].get("nganh_nghe", [])
                    members = entry["data"].get("thanh_vien", [])
                    break
            
            # Chuẩn hóa dữ liệu
            data_lower = {normalize_vietnamese(key): value for key, value in data.items()}
            data_lower["nganh_nghe"] = industries
            data_lower["thanh_vien"] = members

            # Định dạng ngày tháng năm
            try:
                selected_date = datetime.strptime(date_var.get(), "%d/%m/%Y")
                formatted_date = f"ngày {selected_date.day:02d} tháng {selected_date.month:02d} năm {selected_date.year}"
                sort_formatted_date = selected_date.strftime("%d/%m/%Y")  # Định dạng ngày tháng năm dạng 01/04/2025
            except ValueError:
                messagebox.showerror("Lỗi", "Ngày tháng năm không hợp lệ. Vui lòng nhập đúng định dạng DD/MM/YYYY.")
                return

            # Thêm ngày tháng năm vào dữ liệu
            data_lower["ngay_thang_nam"] = formatted_date
            data_lower["sort_ngay_thang_nam"] = sort_formatted_date  # Thêm placeholder mới
            
            # Tính toán và thêm von_dieu_le_bang_chu nếu von_đieu_le tồn tại
            if "von_dieu_le" in data_lower and data_lower["von_dieu_le"]:
                data_lower["von_dieu_le_bang_chu"] = number_to_words(data_lower["von_dieu_le"])
            
            doc_paths = [os.path.join(self.app.templates_dir, t) for t in selected_templates]
            if not all(os.path.exists(dp) for dp in doc_paths):
                missing = [dp for dp in doc_paths if not os.path.exists(dp)]
                messagebox.showerror("Lỗi Template", f"Các file template sau không tồn tại: {', '.join(missing)}")
                return
            
            missing_fields = self.check_template_placeholders(doc_paths, data_lower)
            if missing_fields is None:
                return
            if missing_fields:
                pass
            
            mode = export_mode.get()
            if export_type == "Xem Word":
                self.export_preview(doc_paths, data_lower, mode)
            elif export_type == "Xuất file":
                if file_format.get() == "Word":
                    if mode == "merge":
                        self.export_to_word(doc_paths, data_lower, mode, selected_templates)
                    elif mode == "separate":
                        for doc_path in doc_paths:
                            self.export_to_word([doc_path], data_lower, mode, selected_templates)
                else:
                    self.export_to_pdf(doc_paths, data_lower, mode, selected_templates)

        ttk.Button(popup, text="Xuất File", command=confirm_export, style="primary.TButton").pack(pady=10)

    def merge_documents(self, doc_paths, data_lower):
        if not doc_paths or not isinstance(doc_paths, (list, tuple)):
            raise ValueError("doc_paths phải là một danh sách hợp lệ chứa các đường dẫn template")
        
        temp_files = []
        try:
            # Log dữ liệu đầu vào để kiểm tra
            industries = data_lower.get("nganh_nghe", [])
            members = data_lower.get("thanh_vien", [])
            logging.info(f"Dữ liệu đầu vào - Industries: {len(industries)}, Members: {len(members)}")
            logging.info(f"Doc paths: {doc_paths}")

            for i, doc_path in enumerate(doc_paths):
                if not os.path.exists(doc_path):
                    raise FileNotFoundError(f"File không tồn tại: {doc_path}")
                doc = Document(doc_path)

                # Đặt font mặc định cho toàn bộ tài liệu
                style = doc.styles["Normal"]
                font = style.font
                font.name = "Times New Roman"
                font.size = Pt(14)
                font._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            
                # Xử lý tất cả placeholder trong tài liệu
               # for paragraph in doc.paragraphs:
                # Xử lý tất cả placeholder trong tài liệu
                for paragraph in doc.paragraphs:
                    if "{{ngay_thang_nam}}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{ngay_thang_nam}}", data_lower.get("ngay_thang_nam", ""))
                    if "{{sort_ngay_thang_nam}}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{sort_ngay_thang_nam}}", data_lower.get("sort_ngay_thang_nam", ""))
                
                    # Xử lý {{bang_nganh_nghe}}
                    if "{{bang_nganh_nghe}}" in paragraph.text or "{{ bang_nganh_nghe }}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{bang_nganh_nghe}}", "").replace("{{ bang_nganh_nghe }}", "")
                        p = paragraph._element
                        if industries:
                            table = self.create_industry_table(doc, industries)
                            p.getparent().insert(p.getparent().index(p) + 1, table._element)
                        else:
                            p.getparent().insert(p.getparent().index(p) + 1, doc.add_paragraph("Không có thông tin ngành nghề để hiển thị.")._element)

                    # Xử lý {{bang_hop_thanh_vien}}
                    if "{{bang_hop_thanh_vien}}" in paragraph.text or "{{ bang_hop_thanh_vien }}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{bang_hop_thanh_vien}}", "").replace("{{ bang_hop_thanh_vien }}", "")
                        p = paragraph._element
                        if members:
                            table = self.create_member_table(doc, members)
                            if table:
                                p.getparent().insert(p.getparent().index(p) + 1, table._element)
                        else:
                            p.getparent().insert(p.getparent().index(p) + 1, doc.add_paragraph("Không có thông tin thành viên để hiển thị.")._element)

                    # Xử lý {{bang_thanh_vien}}
                    if "{{bang_thanh_vien}}" in paragraph.text or "{{ bang_thanh_vien }}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{bang_thanh_vien}}", "").replace("{{ bang_thanh_vien }}", "")
                        p = paragraph._element
                        if members:
                            table = self.create_member_info_table(doc, members)
                            if table:
                                p.getparent().insert(p.getparent().index(p) + 1, table._element)
                                logging.info("Đã chèn bang_thanh_vien")
                        else:
                            p.getparent().insert(p.getparent().index(p) + 1, doc.add_paragraph("Không có thông tin thành viên để hiển thị.")._element)

                    # Xử lý {{bang_gop_von}}
                    if "{{bang_gop_von}}" in paragraph.text or "{{ bang_gop_von }}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{bang_gop_von}}", "").replace("{{ bang_gop_von }}", "")
                        p = paragraph._element
                        if members:
                            table = self.create_capital_contribution_table(doc, members)
                            if table:
                                p.getparent().insert(p.getparent().index(p) + 1, table._element)
                                logging.info("Đã chèn bang_gop_von")
                        else:
                            p.getparent().insert(p.getparent().index(p) + 1, doc.add_paragraph("Không có thông tin góp vốn để hiển thị.")._element)

                    # Xử lý {{danh_sach_thanh_vien}}
                    if "{{danh_sach_thanh_vien}}" in paragraph.text or "{{ danh_sach_thanh_vien }}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{danh_sach_thanh_vien}}", "").replace("{{ danh_sach_thanh_vien }}", "")
                        p = paragraph._element
                        if members:
                            table = self.create_member_list_table(doc, members)
                            if table:
                                p.getparent().insert(p.getparent().index(p) + 1, table._element)
                                logging.info("Đã chèn danh_sach_thanh_vien")
                        else:
                            p.getparent().insert(p.getparent().index(p) + 1, doc.add_paragraph("Không có danh sách thành viên để hiển thị.")._element)

                # Lưu file tạm thời
                temp_file = f"temp_{i}.docx"
                temp_files.append(temp_file)
                doc.save(temp_file)
                logging.info(f"Đã lưu file tạm thời: {temp_file}")

                # Tạo render_data cho các placeholder còn lại
                render_data = {k: v for k, v in data_lower.items() if k not in ["nganh_nghe", "thanh_vien"]}
                for idx, member in enumerate(members, start=1):
                    for key, value in member.items():
                        render_data[f"{key}_{idx}"] = value
                # Thêm giá trị giả để tránh cảnh báo
                render_data["bang_nganh_nghe"] = ""
                render_data["bang_hop_thanh_vien"] = ""
                render_data["bang_thanh_vien"] = ""
                render_data["bang_gop_von"] = ""
                render_data["danh_sach_thanh_vien"] = ""

                template = DocxTemplate(temp_file)
                template.render(render_data)
                template.save(temp_file)
                logging.info(f"Đã cập nhật file tạm thời với các placeholder còn lại: {temp_file}")

            if not temp_files:
                return None

            if len(temp_files) == 1:  # Chế độ xuất riêng lẻ
                final_doc = Document(temp_files[0])
                logging.info("Hoàn tất chế độ xuất riêng lẻ")
                return final_doc
            else:  # Chế độ xuất gộp
                master = Document(temp_files[0])
                composer = Composer(master)
                for i, temp_file in enumerate(temp_files[1:], 1):
                    doc = Document(temp_file)
                    if i > 0:
                        add_section_break(composer.doc)
                    composer.append(doc)
                logging.info("Hoàn tất chế độ xuất gộp")
                return master

        except Exception as e:
            logging.error(f"Lỗi khi gộp tài liệu: {str(e)}")
            messagebox.showerror("Lỗi Template", f"Lỗi khi xử lý template:\n{str(e)}")
            return None
        finally:
            for temp_file in temp_files:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
        return None

    def export_preview(self, doc_paths, data_lower, mode):
        if mode == "merge":
            # Tính số lượng thành viên
            so_thanh_vien = len(data_lower.get("thanh_vien", []))
            data_lower["so_thanh_vien"] = str(so_thanh_vien)  # Thêm vào data_lower
            
            merged_doc = self.merge_documents(doc_paths, data_lower)
            if merged_doc:
                preview_window = create_popup(self.app.root, "Xem trước Word", 800, 600)
                text_widget = Text(preview_window, wrap="word")
                text_widget.pack(fill="both", expand=True, padx=10, pady=10)
                for paragraph in merged_doc.paragraphs:
                    text_widget.insert(tk.END, paragraph.text + "\n")
        else:
            messagebox.showinfo("Thông báo", "Chế độ xem trước chỉ hỗ trợ gộp file!")

    def export_to_word(self, doc_paths, data_lower, mode, selected_templates):
        # Tính số lượng thành viên
        so_thanh_vien = len(data_lower.get("thanh_vien", []))
        data_lower["so_thanh_vien"] = str(so_thanh_vien)  # Thêm vào data_lower

        # Lấy tên công ty và template từ dữ liệu
        ma_so_doanh_nghiep = data_lower.get("ma_so_doanh_nghiep", "output")
        
        # Loại bỏ đuôi .docx khỏi tên các template
        template_names = [os.path.splitext(t)[0] for t in selected_templates]
        template_name = "_".join(template_names)  # Gộp tên các template

        if mode == "merge":
            output_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")],
                initialfile=f"{ma_so_doanh_nghiep}_{template_name}_{datetime.now().strftime('%d%m%Y')}.docx"
            )
            if output_path:
                merged_doc = self.merge_documents(doc_paths, data_lower)
                if merged_doc:
                    try:
                        merged_doc.save(output_path)
                        os.startfile(output_path)
                        messagebox.showinfo("Thành công", f"Đã xuất file: {output_path}")
                        logging.info(f"Xuất Word gộp: {output_path}")
                    except Exception as e:
                        logging.error(f"Lỗi khi xuất Word: {str(e)}")
                        messagebox.showerror("Lỗi", f"Không thể xuất Word: {str(e)}")
        elif mode == "separate":
            base_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")],
                initialfile=f"{ma_so_doanh_nghiep}_{template_name}_{datetime.now().strftime('%d%m%Y_%H%M%S')}.docx",
                title="Chọn vị trí lưu file đầu tiên"
            )
            if base_path:
                base_path, ext = os.path.splitext(base_path)
                for i, doc_path in enumerate(doc_paths):
                    single_doc = self.merge_documents([doc_path], data_lower)
                    if single_doc:
                        file_path = f"{base_path}_{i+1}{ext}"
                        try:
                            single_doc.save(file_path)
                            os.startfile(file_path)
                        except Exception as e:
                            logging.error(f"Lỗi khi lưu hoặc mở file {file_path}: {str(e)}")
                            messagebox.showerror("Lỗi", f"Không thể lưu/mở file {file_path}: {str(e)}")
                messagebox.showinfo("Thành công", f"Đã xuất {len(doc_paths)} file Word riêng lẻ!")
                logging.info(f"Xuất {len(doc_paths)} file Word riêng lẻ từ {base_path}")

    def export_to_pdf(self, doc_paths, data_lower, mode):
        # Tính số lượng thành viên
        so_thanh_vien = len(data_lower.get("thanh_vien", []))
        data_lower["so_thanh_vien"] = str(so_thanh_vien)  # Thêm vào data_lower

        output_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")],
                                                    initialfile=f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
        if output_path:
            if mode == "merge":
                merged_doc = self.merge_documents(doc_paths, data_lower)
                if merged_doc:
                    try:
                        temp_docx = "temp_merged.docx"
                        merged_doc.save(temp_docx)
                        convert(temp_docx, output_path)
                        os.remove(temp_docx)
                        os.startfile(output_path)
                        messagebox.showinfo("Thành công", f"Đã xuất file PDF: {output_path}")
                        logging.info(f"Xuất PDF gộp: {output_path}")
                    except Exception as e:
                        logging.error(f"Lỗi khi xuất PDF: {str(e)}")
                        messagebox.showerror("Lỗi", f"Không thể xuất PDF: {str(e)}")
            else:
                base_path, ext = os.path.splitext(output_path)
                for i, doc_path in enumerate(doc_paths):
                    template = DocxTemplate(doc_path)
                    template.render(data_lower)
                    temp_docx = f"temp_{i}.docx"
                    template.save(temp_docx)
                    try:
                        file_path = f"{base_path}_{i+1}.pdf"
                        convert(temp_docx, file_path)
                        os.remove(temp_docx)
                        os.startfile(file_path)
                    except Exception as e:
                        logging.error(f"Lỗi khi chuyển đổi PDF từ {temp_docx}: {str(e)}")
                        messagebox.showerror("Lỗi", f"Không thể chuyển đổi sang PDF: {str(e)}")
                messagebox.showinfo("Thành công", f"Đã xuất {len(doc_paths)} file PDF riêng lẻ!")
                logging.info(f"Xuất {len(doc_paths)} file PDF riêng lẻ từ {base_path}")

    def create_industry_table(self, doc, industries):
        # Tạo bảng với số hàng = 1 (tiêu đề) + số ngành nghề
        table = doc.add_table(rows=1 + len(industries), cols=4)
        table.autofit = False  # Tắt autofit để cố định kích thước

        # Đặt chiều rộng cột
        for cell in table.columns[0].cells:
            cell.width = Cm(1.5)  # Cột 1: STT
        for cell in table.columns[1].cells:
            cell.width = Cm(9.5)  # Cột 2: Tên ngành
        for cell in table.columns[2].cells:
            cell.width = Cm(2.75)  # Cột 3: Mã ngành
        for cell in table.columns[3].cells:
            cell.width = Cm(3.0)  # Cột 4: Ngành nghề kinh doanh chính

        # Thêm viền cho bảng
        table_style = table._tblPr.xpath('./w:tblBorders')[0] if table._tblPr.xpath('./w:tblBorders') else OxmlElement('w:tblBorders')
        for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            table_style.append(border)
        if not table._tblPr.xpath('./w:tblBorders'):
            table._tblPr.append(table_style)

        # Đặt tiêu đề cho các cột
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "STT"
        hdr_cells[1].text = "Tên ngành"
        hdr_cells[2].text = "Mã ngành"
        hdr_cells[3].text = "Ngành nghề kinh doanh chính"

        # Định dạng chữ cho tiêu đề
        for i, cell in enumerate(hdr_cells):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.text = cell.text
            font = run.font
            font.size = Pt(14)
            font.bold = True

        # Thêm dữ liệu ngành nghề
        for idx, industry in enumerate(industries, start=1):
            row_cells = table.rows[idx].cells  # Lấy hàng đã tạo sẵn
            row_cells[0].text = str(idx)
            row_cells[1].text = industry.get("ten_nganh", "")
            row_cells[2].text = industry.get("ma_nganh", "")
            row_cells[3].text = "X" if industry.get("la_nganh_chinh", False) else ""

            # Định dạng chữ cho dữ liệu
            for i, cell in enumerate(row_cells):
                paragraph = cell.paragraphs[0]
                if i in [0, 2, 3]:
                    paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = cell.text
                font = run.font
                font.size = Pt(14)

        # Căn giữa bảng
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        return table

    def create_member_table(self, doc, members):
        """Tạo bảng họp thành viên (bang_hop_thanh_vien)."""
        if not members:
            p = doc.add_paragraph()
            p.text = "Không có thông tin thành viên để hiển thị."
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return None

        # Tạo bảng với 2 cột
        table = doc.add_table(rows=len(members), cols=2)
        table.autofit = True  # bật chế độ tự động điều chỉnh độ rộng

        # Đặt chiều rộng cột
        widths = [5.5, 11.5] 
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width) 

        # Điền dữ liệu vào bảng
        for i, member in enumerate(members, start=1):  
            so_giay_chung_nhan = f"{i:02d}/GCN"  # Tạo số giấy chứng nhận

            # Xác định vai trò dựa trên la_chu_tich
            if member.get("la_chu_tich", False):
                role = "Chủ tịch hội đồng thành viên – Chủ toạ cuộc họp sở hữu"
            else:
                role = "Thành viên góp vốn sở hữu"

            # Cột 1: Tên thành viên
            cell_1 = table.cell(i-1, 0)
            p1 = cell_1.paragraphs[0]
            p1.text = member.get("ho_ten", "")
            p1.style = doc.styles["Normal"]
            p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run1 = p1.runs[0]
            run1.font.name = "Times New Roman"
            run1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run1.font.size = Pt(14)

            # Cột 2: Thông tin
            cell_2 = table.cell(i-1, 1)
            p2 = cell_2.paragraphs[0]
            p2.text = (
                f"– {role} {member.get('von_gop', '')} "
                f"chiếm tỷ lệ {member.get('ty_le_gop', '')}% vốn điều lệ.\n"
                f"Giấy chứng nhận góp vốn số {so_giay_chung_nhan}, cấp ngày {member.get('ngay_gop_von', '')}.\n"
            )
            p2.style = doc.styles["Normal"]
            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run2 = p2.runs[0]
            run2.font.name = "Times New Roman"
            run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run2.font.size = Pt(14)

        return table

    def create_member_info_table(self, doc, members):
        """Tạo bảng thông tin thành viên (bang_thanh_vien) với 7 cột."""
        if not members:
            p = doc.add_paragraph()
            p.text = "Không có thông tin thành viên để hiển thị."
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return None

        # Tạo bảng với tiêu đề và dữ liệu
        table = doc.add_table(rows=1 + len(members), cols=7)
        table.autofit = False

        # Đặt chiều rộng cột
        widths = [1.5, 3.5, 2.75, 1.5, 1.5, 4.25, 3.75]  # 7 cột
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

        # Thêm viền cho bảng
        table_style = table._tblPr.xpath('./w:tblBorders')[0] if table._tblPr.xpath('./w:tblBorders') else OxmlElement('w:tblBorders')
        for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            table_style.append(border)
        if not table._tblPr.xpath('./w:tblBorders'):
            table._tblPr.append(table_style)

        # Thêm tiêu đề
        hdr_cells = table.rows[0].cells
        headers = [
            "STT", 
            "Tên thành viên", 
            "Ngày, tháng, năm sinh đối với thành viên là cá nhân", 
            "Giới tính", 
            "Quốc tịch", 
            "Địa chỉ liên lạc đối với cá nhân, hoặc địa chỉ trụ sở chính đối với tổ chức", 
            "Loại giấy tờ, số, ngày cấp, cơ quan cấp Giấy tờ pháp lý của cá nhân/tổ chức"
        ]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            paragraph = hdr_cells[i].paragraphs[0]
            paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.text = header
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(13)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        # Điền dữ liệu
        for idx, member in enumerate(members, start=1):
            row_cells = table.rows[idx].cells
            row_cells[0].text = str(idx)
            row_cells[1].text = member.get("ho_ten", "")
            row_cells[2].text = member.get("ngay_sinh", "")
            row_cells[3].text = member.get("gioi_tinh", "")
            row_cells[4].text = member.get("quoc_tich", "")
            row_cells[5].text = member.get("dia_chi_lien_lac", "")
            row_cells[6].text = (
                f"{member.get('loai_giay_to', '')}\n"
                f"{member.get('so_cccd', '')}\n"
                f"{member.get('ngay_cap', '')}\n"
                f"{member.get('noi_cap', '')}"
            )

            # Định dạng các ô
            for i, cell in enumerate(row_cells):
                paragraph = cell.paragraphs[0]
                paragraph.alignment = (
                    WD_TABLE_ALIGNMENT.CENTER if i in [0, 3, 4] else WD_ALIGN_PARAGRAPH.CENTER
                )
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = cell.text
                font = run.font
                font.size = Pt(13)
                font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")  # Đảm bảo font cho tiếng Việt

        # Căn giữa bảng
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        return table

    def create_capital_contribution_table(self, doc, members):
        """Tạo bảng góp vốn (bang_gop_von) với 7 cột, gộp cột Vốn góp và thêm cột Ghi chú."""
        logging.info(f"Số lượng thành viên trong create_capital_contribution_table: {len(members)}")
        logging.info(f"Dữ liệu thành viên: {members}")
        
        if not members:
            p = doc.add_paragraph()
            p.text = "Không có danh sách thành viên để hiển thị."
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return None

        # Tạo bảng với 2 hàng tiêu đề (tiêu đề chính gộp, tiêu đề phụ) + số hàng dữ liệu
        table = doc.add_table(rows=2 + len(members), cols=7)
        table.autofit = False

        # Đặt chiều rộng cột
        widths = [1.27, 3.5, 4.41, 1.75, 2.25, 3.5, 1.25]  # 7 cột, bao gồm cột Ghi chú
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

        # Thêm viền cho bảng
        table_style = table._tblPr.xpath('./w:tblBorders')[0] if table._tblPr.xpath('./w:tblBorders') else OxmlElement('w:tblBorders')
        for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            table_style.append(border)
        if not table._tblPr.xpath('./w:tblBorders'):
            table._tblPr.append(table_style)
        
        # Hàng 1: Tiêu đề chính
        main_hdr_cells = table.rows[0].cells
        main_headers = [
            "STT", "Tên thành viên", "Vốn góp", "", "", "Thời hạn góp vốn", "Ghi chú"
        ]
        if len(main_headers) != 7:
            logging.error(f"Số lượng tiêu đề chính không khớp: {len(main_headers)} (yêu cầu 7)")
            raise ValueError("Số lượng tiêu đề chính không khớp với số cột")

        # Gộp hàng 1 và hàng 2 cho cột 1-2 và 6, đồng thời điền tiêu đề chính
        for i in range(len(main_headers)):
            if i in [0, 1, 5, 6]:  # Cột 1-2 (chỉ số 0-1), 6 (chỉ số 5), và 7 (chỉ số 6 - Ghi chú)
                # Gộp ô từ hàng 1 và hàng 2
                cell = main_hdr_cells[i]
                cell.merge(table.rows[1].cells[i])  # Gộp ô hàng 1 và hàng 2
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.add_run(main_headers[i])
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            elif i == 2:  # Cột 3: "Vốn góp"
                cell = main_hdr_cells[2]
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.add_run("Vốn góp")
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                # Gộp cột 3, 4, 5 (chỉ số 2, 3, 4) ở hàng 1
                cell.merge(main_hdr_cells[3])
                cell.merge(main_hdr_cells[4])
            elif i in [3, 4]:
                continue  # Bỏ qua vì đã gộp vào cột 3

        # Hàng 2: Tiêu đề phụ (chỉ hiển thị cho các cột con của "Vốn góp", các cột khác đã gộp)
        sub_hdr_cells = table.rows[1].cells
        for i in range(7):
            if i in [2, 3, 4]:  # Chỉ hiển thị tiêu đề phụ cho cột 3, 4, 5
                sub_headers = ["Phần vốn góp (VNĐ)", "Tỷ lệ (%)", "Loại tài sản, số lượng, giá trị tài sản góp vốn"]
                sub_hdr_cells[i].text = sub_headers[i - 2]
                paragraph = sub_hdr_cells[i].paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = sub_headers[i - 2]
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        # Điền dữ liệu
        for idx, member in enumerate(members, start=1):
            row_cells = table.rows[idx + 1].cells  # +1 vì có 2 hàng tiêu đề
            row_cells[0].text = str(idx)
            row_cells[1].text = member.get("ho_ten", "")
            row_cells[2].text = member.get("von_gop", "")
            row_cells[3].text = member.get("ty_le_gop", "")
            row_cells[4].text = "Đồng Việt Nam"  # Khớp với tài liệu
            row_cells[5].text = member.get("ngay_gop_von", "")
            row_cells[6].text = ""  # Ghi chú để trống

            for i, cell in enumerate(row_cells):
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = cell.text
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        logging.info("Đã chèn bang_gop_von")
        return table

    def create_member_list_table(self, doc, members):
        """Tạo bảng danh sách thành viên (danh_sach_thanh_vien) với 14 cột và gộp cột Vốn góp."""
        logging.info(f"Số lượng thành viên trong create_member_list_table: {len(members)}")
        logging.info(f"Dữ liệu thành viên: {members}")
        
        if not members:
            p = doc.add_paragraph()
            p.text = "Không có danh sách thành viên để hiển thị."
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return None

        # Tạo bảng với 3 hàng tiêu đề (tiêu đề chính gộp, tiêu đề phụ, số thứ tự cột) + số hàng dữ liệu
        table = doc.add_table(rows=3 + len(members), cols=14)
        table.autofit = False

        # Đặt chiều rộng cột
        widths = [1.2, 3.05, 1.62, 1.38, 1.5, 1.5, 3.4, 3.35, 2.75, 1.25, 1.5, 2, 2.5, 1.09]  # Thêm cột Ghi chú
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

        # Thêm viền cho bảng
        table_style = table._tblPr.xpath('./w:tblBorders')[0] if table._tblPr.xpath('./w:tblBorders') else OxmlElement('w:tblBorders')
        for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            table_style.append(border)
        if not table._tblPr.xpath('./w:tblBorders'):
            table._tblPr.append(table_style)
        
        # Hàng 1: Tiêu đề chính
        main_hdr_cells = table.rows[0].cells
        main_headers = [
            "STT", "Tên thành viên", "Ngày, tháng, năm sinh đối với thành viên là cá nhân", "Giới tính", "Quốc tịch", "Dân tộc",
            "Địa chỉ liên lạc đối với thành viên là cá nhân; địa chỉ trụ sở chính đối với thành viên là tổ chức",
            "Loại giấy tờ, số, ngày cấp, cơ quan cấp Giấy tờ pháp lý của cá nhân/tổ chức", "Vốn góp", "", "", "Thời hạn góp vốn",
            "Chữ ký của thành viên", "Ghi chú"
        ]
        if len(main_headers) != 14:
            logging.error(f"Số lượng tiêu đề chính không khớp: {len(main_headers)} (yêu cầu 14)")
            raise ValueError("Số lượng tiêu đề chính không khớp với số cột")

        # Gộp hàng 1 và hàng 2 cho cột 1-8 và 12-14, đồng thời điền tiêu đề chính
        for i in range(len(main_headers)):
            if i in [0, 1, 2, 3, 4, 5, 6, 7, 11, 12, 13]:  # Cột 1-8 (chỉ số 0-7) và 12-14 (chỉ số 11-13)
                # Gộp ô từ hàng 1 và hàng 2
                cell = main_hdr_cells[i]
                cell.merge(table.rows[1].cells[i])  # Gộp ô hàng 1 và hàng 2
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.add_run(main_headers[i])
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            elif i == 8:  # Cột 9: "Vốn góp"
                cell = main_hdr_cells[8]
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.add_run("Vốn góp")
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                # Gộp cột 9, 10, 11 (chỉ số 8, 9, 10) ở hàng 1
                cell.merge(main_hdr_cells[9])
                cell.merge(main_hdr_cells[10])
            elif i in [9, 10]:
                continue  # Bỏ qua vì đã gộp vào cột 9

        # Hàng 2: Tiêu đề phụ (chỉ hiển thị cho các cột con của "Vốn góp", các cột khác đã gộp)
        sub_hdr_cells = table.rows[1].cells
        for i in range(14):
            if i in [8, 9, 10]:  # Chỉ hiển thị tiêu đề phụ cho cột 9, 10, 11
                sub_headers = ["Phần vốn góp (VNĐ)", "Tỷ lệ (%)", "Loại tài sản, số lượng, giá trị tài sản góp vốn"]
                sub_hdr_cells[i].text = sub_headers[i - 8]
                paragraph = sub_hdr_cells[i].paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = sub_headers[i - 8]
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        # Hàng 3: Số thứ tự cột (1, 2, 3, ..., 14)
        num_hdr_cells = table.rows[2].cells
        for i in range(14):
            num_hdr_cells[i].text = str(i + 1)
            paragraph = num_hdr_cells[i].paragraphs[0]
            paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.text = str(i + 1)
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(12)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        # Điền dữ liệu
        for idx, member in enumerate(members, start=1):
            row_cells = table.rows[idx + 2].cells  # +2 vì có 3 hàng tiêu đề
            row_cells[0].text = str(idx)
            row_cells[1].text = member.get("ho_ten", "")
            row_cells[2].text = member.get("ngay_sinh", "")
            row_cells[3].text = member.get("gioi_tinh", "")
            row_cells[4].text = member.get("quoc_tich", "")
            row_cells[5].text = member.get("dan_toc", "")
            row_cells[6].text = member.get("dia_chi_lien_lac", "")
            row_cells[7].text = f"{member.get('loai_giay_to', '')}\n{member.get('so_cccd', '')}\n{member.get('ngay_cap', '')}\n{member.get('noi_cap', '')}"
            row_cells[8].text = member.get("von_gop", "")
            row_cells[9].text = member.get("ty_le_gop", "")
            row_cells[10].text = "Đồng Việt Nam"  # Khớp với tài liệu
            row_cells[11].text = member.get("ngay_gop_von", "")
            row_cells[12].text = ""  # Chữ ký để trống
            row_cells[13].text = ""  # Ghi chú để trống

            for i, cell in enumerate(row_cells):
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = cell.text
                font = run.font
                font.name = "Times New Roman"
                font.size = Pt(12)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        logging.info("Đã chèn danh_sach_thanh_vien")
        return table






